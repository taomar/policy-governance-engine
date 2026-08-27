"""Layout-aware document ingestion into a canonical representation.

Implements the PDF ingestion architecture spec (docs/specs/pdf-ingestion-architecture-v1.md).

WHY THIS EXISTS
---------------
The previous ingestion (``document_extraction._extract_pdf``) iterated page by
page and emitted paragraphs scoped to a single page, falling back to grouping
lines in fixed batches of four when a page had no blank-line breaks. Two
architectural consequences followed, and both are visible in extracted output:

1. A paragraph that crossed a page boundary was split into two clauses, so a
   condition could be separated from its consequence, or a rule from its
   exception. The extractor then saw half a rule and formalized half a rule.
2. Fixed-size line grouping made *chunk boundaries into policy boundaries*,
   which the spec forbids outright (sections 12, 13).

Neither was a bug in one document — per spec section 58, the same failure was
possible in every document, so the fix belongs at the ingestion layer rather
than in a per-document workaround.

HOW IT WORKS
------------
Text is reconstructed from ``page.extract_words()`` rather than
``page.extract_text()``. That is deliberate: words carry geometry (top, x0,
font size), which makes several things decidable that are otherwise guesswork —
which lines belong to a table (bbox containment), which lines are headings
(font size relative to the document's modal body size), where column gutters
fall, and where line boundaries actually are. Because this module *constructs*
the canonical page string itself, every offset it records is exact by
construction rather than recovered by searching, satisfying INVARIANT 4 and
INVARIANT 5.

ARBITRARY UPLOADS
-----------------
Files arrive from users, so nothing here may assume well-formed input. The
degenerate cases that actually occur — scans with no text layer, encrypted
files, two-column layouts, rotated watermark text, right-to-left script,
bordered callouts that ruling-line detection mistakes for tables — are each
handled explicitly and, where they cannot be handled, reported as an
``IngestionDiagnostic`` rather than silently reducing coverage (INVARIANT 9).
A document that yields little text because it is a scan must never be
indistinguishable from a document that genuinely contains little policy.

WHAT IT MUST NEVER DO
---------------------
Insert a word. Cross-page reconstruction concatenates existing fragments in
source order (spec section 10). The only text transformations permitted are
deterministic, code-performed, and recorded on the element:
``line_join_space``, ``line_break_hyphen_join``, ``cross_page_join``,
``table_cell_join``.
"""
from __future__ import annotations

import re
import unicodedata
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from policy_platform.contracts.canonical_document import (
    CanonicalDocument,
    CanonicalElement,
    CanonicalPage,
    ElementType,
    IngestionDiagnostic,
    SourceFragment,
    Transformation,
)
from policy_platform.infrastructure.ingestion.canonical_fidelity import (
    verify_element_text,
)
from policy_platform.infrastructure.ingestion.reading_order import (
    Glyph,
    has_rtl,
    logical_order,
    normalize_presentation_forms,
)

PARSER_NAME = "pdfplumber-layout-v1"
DOCX_PARSER_NAME = "python-docx-layout-v1"

# Words on the same visual line rarely differ in `top` by more than a couple of
# points; anything larger is a genuinely new line. Kept tight because legal PDFs
# often set dense leading, where a loose tolerance merges consecutive lines.
_Y_TOLERANCE = 3.0

# Matches pdfplumber's own default word-splitting behaviour poorly on these
# documents: the default merges adjacent words that lack a real space, yielding
# artifacts like "thispolicy". 1.0 recovers the gap.
_X_TOLERANCE = 1.0

_TERMINAL_PUNCTUATION = ('.', '!', '?', ':', ';', '"', '”', '’', "'", ')', '،', '؛', '.')

_LIST_MARKER_RE = re.compile(
    r"^\s*(?:"
    r"\(\d{1,3}\)"          # (1)
    r"|\d{1,3}[.)]"          # 1.  1)
    r"|\(?[a-zA-Z][.)]"      # a.  (a)
    r"|[ivxIVX]{1,5}[.)]"    # iv.
    r"|[\u2022\u25AA\u25E6\u00B7\u2023\u2043*\-\u2013\u2014]"  # bullets/dashes
    r")\s+"
)

# A new numbered provision is a hard boundary signal: it must never be merged
# onto the tail of the previous page's paragraph (spec section 9). Deliberately
# loose — for boundary detection, over-matching is safe (it only prevents a
# join), whereas under-matching silently welds two provisions together.
_PROVISION_START_RE = re.compile(
    r"^\s*(?:"
    r"(?:Article|ARTICLE|Section|SECTION|Chapter|CHAPTER|Clause|CLAUSE|Part|PART|Annex|ANNEX|Appendix|APPENDIX)\b"
    r"|\u0627\u0644\u0645\u0627\u062F\u0629\b"   # Arabic "the Article"
    r"|\u0627\u0644\u0641\u0635\u0644\b"          # Arabic "the Chapter"
    r"|\u0627\u0644\u0628\u0627\u0628\b"          # Arabic "the Part"
    r"|\d{1,3}(?:\.\d{1,3}){1,4}\.?\s"            # 7.2.1   and   3.4.2.
    r")"
)

# Whether a line is a provision *heading* is a much stricter question than
# whether it *starts* one, and needs its own pattern. "Article (81)" is a
# heading; "Article (81) of the Labor Law during the training term or within"
# is the middle of a sentence that merely happens to begin with a citation.
# Distinguishing them matters twice over: the second form is a real policy
# clause that would vanish from extraction, and it would also be installed as
# the `section` label for every element after it.
_PROVISION_HEADING_RE = re.compile(
    r"^\s*(?:"
    r"Article|ARTICLE|Section|SECTION|Chapter|CHAPTER|Clause|CLAUSE|Part|PART|Annex|ANNEX|Appendix|APPENDIX"
    r"|\u0627\u0644\u0645\u0627\u062F\u0629|\u0627\u0644\u0641\u0635\u0644|\u0627\u0644\u0628\u0627\u0628"
    r")\s*"
    r"\(?\s*[\dIVXLC]{1,6}\s*\)?"                      # 12, (12), IV
    r"(?:\s*(?:bis|BIS|ter|TER|\u0645\u0643\u0631\u0631))?"   # "11 bis"
    r"\s*(?:[:\-\u2013\u2014.]\s*\S.{0,140})?"          # optional ": Definitions"
    r"\s*$"
)

_BOILERPLATE_MIN_PAGE_FRACTION = 0.3
_BOILERPLATE_MIN_PAGES = 3
# Running page furniture is a *band* at the page edge, not a fixed number of
# lines: a footer can be one line or several (a name, a page number and a title
# are three), and a window measured in lines sees only part of a taller one.
# The band is found from the page's own geometry instead — see `_edge_lines`.
# It is capped so that a page whose body happens to be evenly leaded all the way
# to the margin cannot hand its closing sentences to the boilerplate detector.
_BOILERPLATE_EDGE_MAX_LINES = 6
# A gap this many times the page's typical leading separates the furniture band
# from the last line of body text. Matches the factor `_build_blocks` already
# uses to recognise a paragraph break, for the same reason: it is the point at
# which vertical space stops being leading and starts being layout.
_BOILERPLATE_EDGE_GAP_RATIO = 1.6
_DIGIT_RE = re.compile(r"\d+")


@dataclass
class _Line:
    """One visual line of text, with the geometry needed to classify it."""

    text: str
    top: float
    bottom: float
    x0: float
    x1: float
    size: float
    page: int
    start_offset: int = 0
    end_offset: int = 0
    in_table: str | None = None  # table_id when the line sits inside a table bbox
    is_boilerplate: bool = False
    #: The parser's own output for this line, in the order the page paints it,
    #: before any reading-order recovery. Kept so that the transformation stays
    #: auditable: anyone can compare what the page painted against what was
    #: stored. Equal to ``text`` whenever nothing needed recovering.
    visual_text: str = ""

    @property
    def fragment(self) -> SourceFragment:
        return SourceFragment(
            page=self.page,
            start_offset=self.start_offset,
            end_offset=self.end_offset,
            text=self.text,
        )


@dataclass
class _Block:
    """A run of lines that form one logical unit *within a single page*.

    Cross-page merging happens afterwards, on blocks, so that the decision to
    join has both candidates fully assembled rather than guessing line by line.
    """

    element_type: ElementType
    lines: list[_Line] = field(default_factory=list)
    table_id: str | None = None
    table_headers: list[str] | None = None
    cell_text: str | None = None  # set for table rows, whose text is not line-derived

    @property
    def top(self) -> float:
        return min(line.top for line in self.lines) if self.lines else 0.0

    @property
    def page(self) -> int:
        return self.lines[0].page if self.lines else 0


class IngestionError(RuntimeError):
    """Raised when a file cannot be ingested at all (encrypted, corrupt, unsupported)."""


def ingest_document(
    storage_path: str | Path, mime_type: str = "", document_id: str = ""
) -> CanonicalDocument:
    """Ingest any supported upload into the canonical representation."""

    path = Path(storage_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf" or "pdf" in mime_type:
        return ingest_pdf(path, document_id)
    if suffix in (".docx", ".doc") or "wordprocessingml" in mime_type:
        return ingest_docx(path, document_id)
    raise IngestionError(f"unsupported document type: {mime_type or suffix or path.name}")


def ingest_pdf(storage_path: str | Path, document_id: str = "") -> CanonicalDocument:
    """Parse a PDF into the canonical representation.

    Deterministic: the same bytes always produce the same element ids, order,
    and offsets (spec section 45). No model is involved at any point.
    """

    path = Path(storage_path)
    diagnostics: list[IngestionDiagnostic] = []
    pages: list[CanonicalPage] = []
    page_blocks: list[list[_Block]] = []
    raw_pages: list[tuple[int, list[_Line], list[_Block]]] = []

    try:
        pdf = pdfplumber.open(path)
    except Exception as exc:  # encrypted, truncated, not actually a PDF
        raise IngestionError(f"cannot open PDF {path.name}: {exc}") from exc

    with pdf:
        if not pdf.pages:
            raise IngestionError(f"PDF {path.name} contains no pages")
        for page_index, page in enumerate(pdf.pages, start=1):
            try:
                lines, tables, page_diagnostics = _read_page(page, page_index)
            except Exception as exc:  # one bad page must not lose the document
                diagnostics.append(
                    IngestionDiagnostic(
                        code="page_parse_failed",
                        severity="error",
                        page=page_index,
                        detail=str(exc),
                    )
                )
                lines, tables, page_diagnostics = [], [], []
            diagnostics.extend(page_diagnostics)
            raw_pages.append((page_index, lines, tables))

    smallest_body_size, body_size = _body_size_band(
        [line for _, lines, _ in raw_pages for line in lines]
    )
    boilerplate = _detect_boilerplate([lines for _, lines, _ in raw_pages])

    for page_index, lines, table_blocks in raw_pages:
        removed: list[str] = []
        for line in lines:
            if _normalize_line(line.text) in boilerplate and _is_edge_line(line, lines):
                line.is_boilerplate = True
                removed.append(line.text)
        # INVARIANT: raw_text keeps every character the parser produced,
        # including headers and footers. They are removed from the *logical*
        # flow only (spec section 11) — dropping them from raw_text would
        # invalidate every offset already recorded against it.
        raw_text = "\n".join(line.text for line in lines)
        visual_raw_text = "\n".join(line.visual_text or line.text for line in lines)
        pages.append(
            CanonicalPage(
                page=page_index,
                raw_text=raw_text,
                removed_boilerplate=removed,
                visual_order_raw_text=(
                    visual_raw_text if visual_raw_text != raw_text else None
                ),
            )
        )
        page_blocks.append(
            _build_blocks(
                lines, table_blocks, body_size, smallest_body_size=smallest_body_size
            )
        )

    elements = _assemble_elements(page_blocks)

    document = CanonicalDocument(
        document_id=document_id or path.stem,
        page_count=len(pages),
        pages=pages,
        elements=elements,
        parser=PARSER_NAME,
        diagnostics=diagnostics,
    )
    _append_document_diagnostics(document)
    return document


def _append_document_diagnostics(document: CanonicalDocument) -> None:
    """Report degradations that would otherwise look like a sparse document."""

    empty_pages = [page.page for page in document.pages if not page.raw_text.strip()]
    if empty_pages and len(empty_pages) >= max(1, len(document.pages) // 2):
        document.diagnostics.append(
            IngestionDiagnostic(
                code="no_text_layer",
                severity="error",
                detail=(
                    f"{len(empty_pages)} of {len(document.pages)} pages contain no extractable "
                    "text. The file is most likely a scan and needs OCR before it can be used "
                    "for policy extraction."
                ),
            )
        )
    elif empty_pages:
        document.diagnostics.append(
            IngestionDiagnostic(
                code="empty_pages",
                severity="warning",
                detail=f"pages with no extractable text: {empty_pages[:20]}",
            )
        )

    if not document.elements:
        document.diagnostics.append(
            IngestionDiagnostic(
                code="no_elements",
                severity="error",
                detail="ingestion produced no logical elements",
            )
        )
        return

    raw_chars = sum(
        len(_strip_whitespace(page.raw_text)) for page in document.pages
    )
    element_chars = sum(len(_strip_whitespace(element.text)) for element in document.elements)
    if raw_chars:
        ratio = element_chars / raw_chars
        if ratio < 0.85:
            document.diagnostics.append(
                IngestionDiagnostic(
                    code="low_coverage",
                    severity="warning",
                    detail=(
                        f"only {ratio:.1%} of source characters reached a logical element; "
                        "content may be trapped in unparsed regions"
                    ),
                )
            )

    failures = document.verify_fragments()
    if failures:
        document.diagnostics.append(
            IngestionDiagnostic(
                code="fragment_offsets_unresolvable",
                severity="error",
                detail=f"{len(failures)} source fragments do not resolve; first: {failures[0]}",
            )
        )

    # The next link in the same chain. `verify_fragments` above proves a
    # fragment resolves to its offsets; this proves the element's stored text is
    # those fragments joined as the element itself declares. Between them sits
    # the only step where canonical text is *written*, and it was unchecked.
    #
    # It matters because the checks on either side stay green while it is wrong.
    # The fragments still resolve, and the extraction agent still copies
    # faithfully from what it was shown -- so `verify_verbatim` compares two
    # copies of the same corruption and passes. Published v1 carries seven
    # passages that reached a customer-facing citation exactly that way.
    #
    # A warning rather than an error: unlike an unresolvable offset, a rebuilt
    # mismatch can be a transformation this element declared and the rebuild
    # does not model, and refusing an ingest on that would be the check
    # asserting more than it establishes.
    fidelity = verify_element_text(document)
    if fidelity.failures:
        document.diagnostics.append(
            IngestionDiagnostic(
                code="element_text_not_rebuilt_from_fragments",
                severity="warning",
                detail=(
                    f"{len(fidelity.failures)} of {fidelity.checked} elements carry text that is "
                    f"not their recorded fragments joined as declared; "
                    f"first: {fidelity.failures[0]}"
                ),
            )
        )
    elif document.elements and fidelity.checked == 0:
        # The distinction the report exists to preserve: no failures because
        # nothing could be checked is not the same as no failures because
        # everything passed, and only one of them is good news.
        document.diagnostics.append(
            IngestionDiagnostic(
                code="element_text_unprovable",
                severity="info",
                detail=(
                    f"none of the {len(document.elements)} elements could be checked against "
                    "their fragments, so canonical text is unverified rather than verified"
                ),
            )
        )

    rtl_chars = sum(
        1
        for element in document.elements
        for char in element.text
        if _is_rtl(char)
    )
    total_letters = sum(
        1 for element in document.elements for char in element.text if char.isalpha()
    )
    if total_letters and rtl_chars / total_letters > 0.2:
        recovered_pages = [
            page.page for page in document.pages if page.visual_order_raw_text is not None
        ]
        document.diagnostics.append(
            IngestionDiagnostic(
                code="rtl_script_detected",
                severity="info",
                detail=(
                    "right-to-left script detected. Characters are ordered within each "
                    "directional run and the runs ordered by their own direction, so the text "
                    "is stored in reading order rather than paint order. Left-to-right runs "
                    "embedded in it, such as numbers and Latin terms, keep their own order. "
                    f"Pages whose paint order differed from reading order: {len(recovered_pages)}. "
                    "The parser's unmodified output for those pages is retained on each page as "
                    "visual_order_raw_text."
                ),
            )
        )


def _strip_whitespace(text: str) -> str:
    return "".join(text.split())


def _is_rtl(char: str) -> bool:
    return unicodedata.bidirectional(char) in ("R", "AL")


def _read_page(page, page_index: int) -> tuple[list[_Line], list[_Block], list[IngestionDiagnostic]]:
    """Reconstruct one page's lines (with exact offsets) and its tables.

    Lines are built and offset *before* tables are resolved, because a table
    row's provenance is the span of real source lines it covers. Deriving row
    offsets any other way would mean inventing positions, and a source position
    that does not resolve is worse than no position at all — it makes an
    unverifiable extraction look verified.
    """

    diagnostics: list[IngestionDiagnostic] = []
    try:
        words = page.extract_words(
            x_tolerance=_X_TOLERANCE, extra_attrs=["size"], return_chars=True
        )
    except Exception:  # malformed font metrics; retry without the extra attribute
        words = page.extract_words(x_tolerance=_X_TOLERANCE, return_chars=True)

    # Rotated text is almost always a watermark, sidebar label, or stamp. Left
    # in the flow it interleaves with body text at arbitrary positions and
    # corrupts reading order, so it is excluded and reported rather than
    # silently mixed in.
    upright_words = [word for word in words if word.get("upright", True)]
    rotated_count = len(words) - len(upright_words)
    if rotated_count:
        diagnostics.append(
            IngestionDiagnostic(
                code="rotated_text_excluded",
                severity="info",
                page=page_index,
                detail=f"{rotated_count} rotated words excluded from the logical flow",
            )
        )

    if not upright_words:
        return [], [], diagnostics

    lines = _group_words_into_lines(upright_words, page_index)

    columns = _detect_columns(lines, float(getattr(page, "width", 0.0) or 0.0))
    if columns:
        diagnostics.append(
            IngestionDiagnostic(
                code="multi_column_layout",
                severity="info",
                page=page_index,
                detail=f"{len(columns)} columns detected; lines ordered per column",
            )
        )
        lines = _order_by_column(lines, columns)

    offset = 0
    for index, line in enumerate(lines):
        line.start_offset = offset
        line.end_offset = offset + len(line.text)
        offset = line.end_offset + (1 if index < len(lines) - 1 else 0)

    table_regions: list[tuple[str, tuple[float, float, float, float]]] = []
    table_blocks: list[_Block] = []
    try:
        found = page.find_tables()
    except Exception:  # pdfplumber can throw on degenerate ruling geometry
        found = []
    for table_index, table in enumerate(found, start=1):
        table_id = f"p{page_index}-t{table_index}"
        blocks, table_diagnostics = _table_to_blocks(table, table_id, page_index, lines)
        if not blocks:
            # Not a genuine table (see _is_genuine_table). Leaving the region
            # unregistered is the whole point: the lines then flow through
            # normal paragraph grouping and get reassembled correctly.
            continue
        diagnostics.extend(table_diagnostics)
        table_regions.append((table_id, table.bbox))
        table_blocks.extend(blocks)

    for line in lines:
        line.in_table = _containing_table(line, table_regions)

    return lines, table_blocks, diagnostics


# A gutter narrower than this is word spacing, not a column separator.
_MIN_GUTTER_POINTS = 16.0
# Fraction of lines that may straddle a candidate gutter before it is rejected.
_MAX_GUTTER_CROSSING = 0.08
# Each column must carry at least this share of the page's lines.
_MIN_COLUMN_SHARE = 0.15


def _detect_columns(
    lines: list[_Line], page_width: float
) -> list[tuple[float, float]]:
    """Find vertical gutters that separate genuine text columns.

    Two-column layouts are common in regulations and handbooks, and grouping
    words by vertical position alone welds the left and right column of the
    same visual line into one string — silently producing text that reads as
    nonsense while looking perfectly well-formed. That failure is invisible
    downstream, so it has to be caught here.

    A candidate gutter is an interior band of x with no text in it. It is only
    accepted when almost no line crosses it (a real gutter is respected by
    every line) and both sides carry a meaningful share of the page's lines
    (which rejects the ragged right margin of ordinary prose).
    """

    if page_width <= 0 or len(lines) < 8:
        return []

    width = int(page_width) + 1
    occupied = bytearray(width)
    for line in lines:
        start = max(0, int(line.x0))
        end = min(width - 1, int(line.x1))
        for x in range(start, end + 1):
            occupied[x] = 1

    gaps: list[tuple[int, int]] = []
    run_start: int | None = None
    for x in range(width):
        if not occupied[x]:
            if run_start is None:
                run_start = x
        elif run_start is not None:
            gaps.append((run_start, x))
            run_start = None

    total = len(lines)
    accepted: list[tuple[int, int]] = []
    for start, end in gaps:
        if end - start < _MIN_GUTTER_POINTS:
            continue
        if start < page_width * 0.15 or end > page_width * 0.85:
            continue  # page margins, not a gutter
        centre = (start + end) / 2
        crossing = sum(1 for line in lines if line.x0 < centre < line.x1)
        if crossing > total * _MAX_GUTTER_CROSSING:
            continue
        left = sum(1 for line in lines if line.x1 <= centre)
        right = sum(1 for line in lines if line.x0 >= centre)
        if left < total * _MIN_COLUMN_SHARE or right < total * _MIN_COLUMN_SHARE:
            continue
        accepted.append((start, end))

    if not accepted:
        return []

    boundaries = [0.0]
    for start, end in accepted:
        boundaries.append((start + end) / 2)
    boundaries.append(page_width)
    return [(boundaries[i], boundaries[i + 1]) for i in range(len(boundaries) - 1)]


def _order_by_column(
    lines: list[_Line], columns: list[tuple[float, float]]
) -> list[_Line]:
    """Read each column top-to-bottom before moving to the next."""

    buckets: list[list[_Line]] = [[] for _ in columns]
    for line in lines:
        centre = (line.x0 + line.x1) / 2
        index = 0
        for candidate, (start, end) in enumerate(columns):
            if start <= centre < end:
                index = candidate
                break
        else:
            index = len(columns) - 1
        buckets[index].append(line)
    ordered: list[_Line] = []
    for bucket in buckets:
        ordered.extend(sorted(bucket, key=lambda item: (item.top, item.x0)))
    return ordered


def _group_words_into_lines(words: list[dict], page_index: int) -> list[_Line]:
    if not words:
        return []
    ordered = sorted(words, key=lambda w: (round(w["top"], 1), w["x0"]))
    lines: list[_Line] = []
    current: list[dict] = []
    current_top: float | None = None

    for word in ordered:
        if current_top is None or abs(word["top"] - current_top) <= _Y_TOLERANCE:
            if current_top is None:
                current_top = word["top"]
            current.append(word)
        else:
            lines.append(_line_from_words(current, page_index))
            current = [word]
            current_top = word["top"]
    if current:
        lines.append(_line_from_words(current, page_index))
    return lines


def _text_from_words(words: list[dict]) -> tuple[str, str]:
    """Return ``(logical_text, visual_text)`` for one visual line of words.

    A PDF records where each glyph was *painted*, not the order in which the
    words were written. For text whose script runs left to right the two
    coincide, so left-to-right pages take the original path unchanged and come
    out byte for byte identical. Where a run of the line runs right to left they
    do not coincide, and ordering by horizontal position alone stores the line
    backwards — text that is not the document's words.

    Recovery is delegated to :mod:`reading_order`, which orders characters by
    coordinate *within each directional run* and orders the runs themselves by
    direction. It never reverses a string: reversal would also reverse the
    digits of any number embedded in the run, turning 50% into 05%.
    """

    ordered = sorted(words, key=lambda word: word["x0"])
    visual = " ".join(word["text"] for word in ordered)
    if not has_rtl(visual):
        return visual, visual

    glyphs: list[Glyph] = []
    for index, word in enumerate(ordered):
        chars = word.get("chars") or []
        if chars:
            glyphs.extend(
                Glyph(
                    text=char["text"],
                    x0=float(char["x0"]),
                    x1=float(char["x1"]),
                    group=index,
                )
                for char in chars
            )
            continue
        # Degraded path: the parser gave no per-character geometry for this
        # word. Its characters are still known to be painted left to right
        # across its own box, so spacing them evenly across that box reproduces
        # their relative visual positions, which is all the ordering needs.
        text = word["text"]
        x0, x1 = float(word["x0"]), float(word["x1"])
        step = (x1 - x0) / len(text) if text else 0.0
        glyphs.extend(
            Glyph(text=char, x0=x0 + step * offset, x1=x0 + step * (offset + 1), group=index)
            for offset, char in enumerate(text)
        )

    pieces: list[str] = []
    previous_group: int | None = None
    for glyph in logical_order(glyphs):
        if previous_group is not None and glyph.group != previous_group:
            pieces.append(" ")
        pieces.append(glyph.text)
        previous_group = glyph.group
    return normalize_presentation_forms("".join(pieces)), visual


def _line_from_words(words: list[dict], page_index: int) -> _Line:
    ordered = sorted(words, key=lambda w: w["x0"])
    sizes = [float(w.get("size", 0.0) or 0.0) for w in ordered]
    text, visual_text = _text_from_words(ordered)
    return _Line(
        text=text,
        visual_text=visual_text,
        top=min(float(w["top"]) for w in ordered),
        bottom=max(float(w["bottom"]) for w in ordered),
        x0=min(float(w["x0"]) for w in ordered),
        x1=max(float(w["x1"]) for w in ordered),
        size=max(sizes) if sizes else 0.0,
        page=page_index,
    )


def _containing_table(
    line: _Line, regions: list[tuple[str, tuple[float, float, float, float]]]
) -> str | None:
    centre_y = (line.top + line.bottom) / 2
    centre_x = (line.x0 + line.x1) / 2
    for table_id, (x0, top, x1, bottom) in regions:
        if top <= centre_y <= bottom and x0 <= centre_x <= x1:
            return table_id
    return None


def _banner_columns(row_objects: list) -> dict[int, int]:
    """Row-0 cells that span several row-1 cells, mapped to how many they cover.

    A two-row table header — a merged banner over the sub-labels that divide it —
    is emitted by ``table.extract()`` as a row 0 with empty strings beside the
    banner and a row 1 that looks like data. Read from the strings alone the two
    rows are indistinguishable from a legitimate stub crosstab (``['', 'Q1',
    'Q2']`` over ``['North', '10', '20']``), where row 0 *is* the whole header
    and row 1 *is* data. That is why this was left alone for a while: any
    string-only rule had to be balanced between the two, and tuning that balance
    to the corpora in hand is what constraint 1 forbids.

    Geometry settles it without appealing to any document's words. pdfplumber
    gives each row a cell per column position, ``None`` where no cell boundary
    starts. On the reproduction case (GMU staff handbook, page 30) row 0 reads
    ``[(56,80), (80,156), (156,508), None, None, None, None]`` and row 1 reads
    ``[None, None, (156,209), (209,265), (265,380), (380,448), (448,508)]``:
    the banner at column 2 covers x 156–508, which row 1 divides into five. In
    an ordinary table each row-0 cell covers exactly one row-1 cell, so nothing
    is flagged.

    Measured over every PDF in the corpus: 84 tables with two or more rows, 4
    flagged, and all 4 are that same GMU table in duplicate copies of the file.
    No other table in any document trips it.

    Returns ``{column_index: covered_count}`` for the spanning cells, empty when
    the header occupies one row — which is the ordinary case.
    """

    if len(row_objects) < 2:
        return {}
    row0 = row_objects[0].cells
    row1 = [cell for cell in row_objects[1].cells if cell is not None]
    if not row1:
        return {}

    spanning: dict[int, int] = {}
    for column_index, cell in enumerate(row0):
        if cell is None:
            continue
        left, right = cell[0], cell[2]
        # A tolerance of one point: these are floats from a rendered page, and
        # a sub-label's edge is drawn on the banner's edge, not near it.
        covered = sum(1 for other in row1 if other[0] >= left - 1 and other[2] <= right + 1)
        if covered > 1:
            spanning[column_index] = covered
    return spanning


def _join_header_rows(banner: str, sub_label: str) -> str:
    """One column's label when the header occupies two rows.

    Both halves are kept verbatim and the separator is structural, the same
    contract the row text itself keeps with ``" | "``: nothing here invents a
    word that is not in the document. A column under the banner reads
    ``Itemization · Salary Range``; a column beside it, where only one row
    carries a label, reads that label alone rather than gaining a stray
    separator.
    """

    banner, sub_label = banner.strip(), sub_label.strip()
    if banner and sub_label:
        # Already-identical halves are one label written twice, not two.
        return banner if banner == sub_label else f"{banner} · {sub_label}"
    return banner or sub_label


def _table_to_blocks(
    table, table_id: str, page_index: int, lines: list[_Line]
) -> tuple[list[_Block], list[IngestionDiagnostic]]:
    """Preserve table structure rather than flattening it into prose.

    Spec section 31: approval matrices, penalty schedules and expense limits
    live in tables, and a row loses its meaning when its headers are discarded.
    Headers are carried alongside the row rather than woven into a sentence,
    because writing "Tier: 2; Limit: 5000" as prose would be text that does not
    appear anywhere in the document.

    Each row's provenance is the set of real source lines its bounding box
    covers, so the canonical row text is a *recorded transformation* of genuine
    source spans (``table_cell_join``) rather than a free-floating string.

    ``find_tables`` is a per-page API, so what arrives here is one page's grid,
    never a whole table. Row 0 of a page's grid is the table's header row only
    on the page the table starts on; on a continuation page it is content, and
    on a page whose table opens with a banner it is neither. Deciding those
    three cases by position alone -- which is what "row 0 is the header" does --
    both mislabels the rows beneath it and discards the row itself. See
    `_row_states_column_labels`.
    """

    diagnostics: list[IngestionDiagnostic] = []
    try:
        rows = table.extract()
        row_objects = list(table.rows)
    except Exception:  # pragma: no cover
        return [], diagnostics
    if not rows:
        return [], diagnostics

    if not _is_genuine_table(rows):
        return [], diagnostics

    has_headers, header_diagnostic = _column_labels_for(
        rows, table_id=table_id, page=page_index
    )

    # A header split across two rows is still one header. When row 0 carries a
    # banner over columns that row 1 sub-divides, both rows are the header: read
    # row 1 as data and it becomes a phantom provision whose text is the
    # sub-labels themselves, filed under headers that are empty for exactly the
    # columns it names.
    banner = _banner_columns(row_objects) if has_headers and len(rows) > 1 else {}

    if banner:
        headers = [
            _join_header_rows(
                _cell_in_reading_order(table, row_objects, 0, column_index, (rows[0][column_index] or "").strip()),
                _cell_in_reading_order(table, row_objects, 1, column_index, (rows[1][column_index] or "").strip()),
            )
            for column_index in range(len(rows[0]))
        ]
    else:
        headers = (
            [
                _cell_in_reading_order(table, row_objects, 0, column_index, (cell or "").strip())
                for column_index, cell in enumerate(rows[0])
            ]
            if has_headers
            else []
        )

    if header_diagnostic is not None:
        diagnostics.append(header_diagnostic)

    if banner:
        widest = max(banner.values())
        diagnostics.append(
            IngestionDiagnostic(
                code="table_header_spans_multiple_rows",
                severity="info",
                page=page_index,
                detail=(
                    f"table {table_id}: row 0 carries {len(banner)} banner cell(s) over columns "
                    f"row 1 sub-divides (widest covers {widest}); rows 0-1 read as one header"
                ),
            )
        )

    if banner:
        data_rows = list(enumerate(rows))[2:]
    elif has_headers:
        data_rows = list(enumerate(rows))[1:]
    else:
        data_rows = list(enumerate(rows))

    blocks: list[_Block] = []
    for row_index, row in data_rows:
        cells = [
            _cell_in_reading_order(table, row_objects, row_index, column_index, (cell or "").strip())
            for column_index, cell in enumerate(row)
        ]
        if not any(cells):
            continue
        # A pipe-joined cell list keeps every cell value verbatim and adds no
        # words of its own; the separator is structural, not editorial.
        cell_text = " | ".join(cells)

        bbox = None
        if row_index < len(row_objects):
            bbox = getattr(row_objects[row_index], "bbox", None)
        row_lines = _lines_in_bbox(lines, bbox) if bbox else []
        if not row_lines:
            # Coarser but still true: fall back to the whole table's span rather
            # than fabricating a position.
            row_lines = _lines_in_bbox(lines, table.bbox)

        if not row_lines:
            continue

        blocks.append(
            _Block(
                element_type="table_row",
                lines=row_lines,
                table_id=table_id,
                table_headers=headers if has_headers else None,
                cell_text=cell_text,
            )
        )
    return blocks, diagnostics


#: The longest a cell may be and still be read as a column label. A header names
#: a column; content states something about one. The bound is generous for a
#: label rather than fitted to any document: across the two corpora held here
#: the longest genuine header cell is 24 characters ("Abbreviations / Acronyms")
#: and the shortest content cell that a page-0 rule mislabelled as a header is
#: 58 ("Eating in an unauthorized place or at an unauthorized time."), so the
#: bound sits between them with room on both sides rather than against either.
_MAX_HEADER_CELL_CHARS = 40


def _row_states_column_labels(rows: list[list[str | None]]) -> tuple[bool, str]:
    """Whether row 0 of this grid states column labels, and if not, why not.

    `find_tables` runs per page, so "row 0" means "the first row of this page's
    grid" and nothing more. Treating it as the header row unconditionally is
    wrong in three ways that were all measured on real documents:

    * on a **continuation page** row 0 is an ordinary data row. Consuming it
      deletes a provision -- the row is emitted as no element at all -- and
      hands its text to every row beneath it as their column labels, so a
      reviewer is shown one offence's wording as the heading of another;
    * on a page whose table opens with a **merged banner**, row 0 is a single
      spanning cell naming a family of provisions, not a set of labels;
    * on a page whose table has a **two-row header**, row 0 is half of one.

    The test is on the form of the cells, never on their content, so it keys on
    nothing document-specific -- no heading text, no numbering scheme, no
    layout. A row states column labels when several cells are filled, each is
    short enough to be a label, and none of them recurs further down the grid,
    because labels are distinct by construction and values repeat.

    The failure directions are deliberately asymmetric. Judging a real header
    to be content emits it as a row: the text survives and a diagnostic says
    the headers are unknown. Judging content to be a header deletes it. Only
    one of those loses a provision, so the test is written to be reluctant.
    """

    if not rows:
        return False, "the grid has no rows"

    first = [(cell or "").strip() for cell in rows[0]]
    filled = [cell for cell in first if cell]

    if len(filled) < 2:
        return False, (
            f"row 0 fills {len(filled)} of {len(first)} cells, which is a banner "
            "or a spanning title rather than a set of labels"
        )

    longest = max(filled, key=len)
    if len(longest) > _MAX_HEADER_CELL_CHARS:
        return False, (
            f"row 0 carries a {len(longest)}-character cell, longer than a "
            f"column label ({_MAX_HEADER_CELL_CHARS}), so it states content"
        )

    below = {
        (cell or "").strip()
        for row in rows[1:]
        for cell in row
        if (cell or "").strip()
    }
    recurring = sorted(set(filled) & below)
    if recurring:
        return False, (
            f"{len(recurring)} of row 0's cells recur further down the grid, "
            "so they are values rather than labels"
        )

    return True, ""


def _column_labels_for(
    rows: list[list[str | None]], *, table_id: str, page: int | None
) -> tuple[bool, IngestionDiagnostic | None]:
    """Decide whether row 0 states column labels, and report it when it does not.

    Both parsers call this. The decision itself is `_row_states_column_labels`,
    and the wording of the report lives here, so neither can drift from the
    other: the standard was established once, on a PDF, and then held on one
    path only while the other went on consuming row 0 on an assumption. A
    reviewer must be able to discover that a table's header was not evidenced,
    and why, whichever parser read the document.

    Only the report is shared. How a header row's cells are *read* is properly
    per-parser -- a PDF cell has geometry and may need its right-to-left run
    recovered, a DOCX cell arrives as characters already -- so this returns the
    verdict and leaves each caller to take the cells its own way.
    """

    has_headers, why_not = _row_states_column_labels(rows)
    if has_headers:
        return True, None

    # Reported at `warning`, not `info`. A reader uses `info` to decide what to
    # skip, and this is not the benign common case wearing an unusual hat: every
    # row of this grid reaches extraction with no column labels at all, and a
    # reviewer judging coverage cannot discover that from the rows themselves,
    # which look complete.
    return False, IngestionDiagnostic(
        code="table_header_row_not_identified",
        severity="warning",
        page=page,
        detail=(
            f"table {table_id}: no row states column labels ({why_not}); "
            f"{len(rows)} row(s) carry no headers and row 0 is read as content"
        ),
    )


def _cell_in_reading_order(
    table, row_objects: list, row_index: int, column_index: int, extracted: str
) -> str:
    """Return one table cell's text in reading order.

    ``table.extract`` hands back strings with no geometry attached, so a cell
    holding a right-to-left run arrives in paint order and cannot be repaired
    from the string alone. The cell's own bounding box is known, though, so the
    characters inside it can be read again with their coordinates and put
    through the same recovery the body text uses.

    Cells with no right-to-left character return the parser's original string
    untouched, so ordinary tables are unaffected.
    """

    if not has_rtl(extracted):
        return extracted
    try:
        bbox = row_objects[row_index].cells[column_index]
        if not bbox:
            return extracted
        words = table.page.crop(bbox).extract_words(
            x_tolerance=_X_TOLERANCE, return_chars=True
        )
    except Exception:  # pragma: no cover - degenerate cell geometry
        return extracted
    if not words:
        return extracted

    lines: dict[float, list[dict]] = {}
    for word in words:
        if not word.get("upright", True):
            continue
        key = round(float(word["top"]) / _Y_TOLERANCE)
        lines.setdefault(key, []).append(word)
    rendered = [
        _text_from_words(sorted(group, key=lambda word: word["x0"]))[0]
        for _, group in sorted(lines.items())
    ]
    recovered = " ".join(part for part in rendered if part).strip()
    return recovered or extracted


def _is_genuine_table(rows: list[list[str | None]]) -> bool:
    """Reject bordered text boxes that pdfplumber reports as one-column tables.

    ``find_tables`` works from ruling lines, so any framed callout, shaded
    note, or boxed article body registers as a table. Accepting those is
    actively harmful: each visual line becomes its own "row", which shreds a
    single paragraph into fragments and splits sentences — the exact failure
    spec section 13 prohibits. It is strictly worse than not detecting a table
    at all, because the fragments are then also excluded from paragraph
    reassembly.

    A real policy table (approval matrix, penalty schedule, expense tier) has
    values lined up in at least two columns across at least two rows. A framed
    paragraph has content in exactly one column. That structural difference is
    decidable without heuristics about the text itself.
    """

    if len(rows) < 2:
        return False
    column_fill: Counter[int] = Counter()
    for row in rows:
        for index, cell in enumerate(row):
            if (cell or "").strip():
                column_fill[index] += 1
    populated_columns = [index for index, count in column_fill.items() if count >= 2]
    return len(populated_columns) >= 2


def _lines_in_bbox(
    lines: list[_Line], bbox: tuple[float, float, float, float] | None
) -> list[_Line]:
    if not bbox:
        return []
    x0, top, x1, bottom = bbox
    inside: list[_Line] = []
    for line in lines:
        centre_y = (line.top + line.bottom) / 2
        centre_x = (line.x0 + line.x1) / 2
        if top <= centre_y <= bottom and x0 <= centre_x <= x1:
            inside.append(line)
    return inside


def _modal_body_size(sizes: Counter[float]) -> float:
    """The most common font size is body text; anything larger is structural."""

    if not sizes:
        return 0.0
    return sizes.most_common(1)[0][0]


# The width at which a line counts as running the full column, and the share of
# a size's lines that must do so before it counts as running text. Both are read
# off the document rather than assumed, so a narrow column and a wide one are
# judged by their own measure.
_MEASURE_PERCENTILE = 0.9
_FILLS_THE_MEASURE_RATIO = 0.9

# What it takes for a font size to count as one of the document's text classes:
# it must carry a real share of the lines, and those lines must wrap.
_BODY_CLASS_MIN_SHARE = 0.1
_BODY_CLASS_MIN_WRAP = 0.15


def _column_measure(lines: list[_Line]) -> float:
    """How wide the running text runs, taken from the lines themselves."""

    widths = sorted(
        line.x1 - line.x0
        for line in lines
        if line.text.strip() and not line.is_boilerplate
    )
    if not widths:
        return 0.0
    return widths[min(int(len(widths) * _MEASURE_PERCENTILE), len(widths) - 1)]


def _body_size_band(lines: list[_Line]) -> tuple[float, float]:
    """The range of sizes a document sets its running text in.

    A document need not have one typical size. A bilingual document sets each
    language in its own face and the two are co-equal — neither is a heading in
    the other's terms. Asking such a document for a single modal size forces a
    wrong answer: the modal describes one class and slanders the other, and
    every line of the larger class then measures as "set larger than the body"
    and is read as a heading. On the handbook this leaves 459 of 940 lines
    satisfying the heading test before any other evidence is weighed.

    So the body is taken as a band rather than a point. Sizes above it are
    structural, sizes below it are furniture, and a size inside it is body text
    whichever class it belongs to.

    A size joins the band on two structural tests, and it needs both:

    - it carries a real share of the document's lines, and
    - the lines set in it *wrap* — they run the full column measure and continue
      on the next line.

    Frequency alone is not enough, because a document that is mostly headings
    would nominate one and the band would swallow its own structure. Wrapping is
    what makes a class running text: prose runs to the edge of the column, while
    a heading stops where its words stop. Neither test knows anything about
    which alphabet is in use, how the document is numbered, or what it is about.

    A document with a single text class yields a band of zero width and behaves
    exactly as the modal did, so the ordinary case is unchanged. If no size
    passes both tests — a very short document, or one with no running text at
    all — the modal is used, which is the answer given today.
    """

    live = [line for line in lines if line.text.strip()]
    if not live:
        return 0.0, 0.0

    counts: Counter[float] = Counter(round(line.size, 1) for line in live)
    modal = _modal_body_size(counts)

    measure = _column_measure(live)
    if measure <= 0:
        return modal, modal

    wrapping: Counter[float] = Counter(
        round(line.size, 1)
        for line in live
        if (line.x1 - line.x0) >= measure * _FILLS_THE_MEASURE_RATIO
    )

    band = [
        size
        for size, count in counts.items()
        if count >= len(live) * _BODY_CLASS_MIN_SHARE
        and wrapping[size] >= count * _BODY_CLASS_MIN_WRAP
    ]
    if not band:
        return modal, modal
    return min(band), max(band)


def _normalize_line(text: str) -> str:
    return _DIGIT_RE.sub("#", " ".join(text.strip().split())).lower()


def _edge_lines(lines: list[_Line]) -> list[_Line]:
    """The contiguous band of lines at each page edge, whatever its height.

    Running furniture is set apart from the body by vertical space, not by line
    count. A footer carrying a name, a page number and a document title is three
    lines; a plain page number is one. Taking a fixed number of lines from each
    edge sees all of the short footer and only the last line of the tall one,
    which leaves the rest to be classified as content — and a repeated line that
    survives classification becomes a heading, and then the section label for
    everything that follows it.

    So the band is read off the page instead: walk inward from the edge while
    consecutive lines sit within normal leading of one another, and stop at the
    first gap wide enough to be layout rather than leading. On a page with no
    furniture this simply yields the outermost line or two, which no repetition
    threshold will accept, so nothing changes for such a page.

    The walk is capped, because a page whose body runs evenly to the margin has
    no such gap, and an uncapped walk would offer most of the page to the
    boilerplate detector.
    """

    non_empty = [line for line in lines if line.text.strip()]
    if not non_empty:
        return []

    leading = _typical_leading(non_empty)
    limit = leading * _BOILERPLATE_EDGE_GAP_RATIO if leading > 0 else 0.0

    def separation(earlier: _Line, later: _Line) -> float:
        """Vertical space between two lines, in the units `_typical_leading` uses.

        Negative when the boxes overlap, which is common inside a furniture band
        where a page number set slightly larger shares a baseline with the text
        beside it. Overlap is contiguity, so only a positive gap can end a band.
        """

        return later.top - earlier.bottom

    def band(*, from_top: bool) -> list[_Line]:
        ordered = non_empty if from_top else list(reversed(non_empty))
        collected = [ordered[0]]
        set_apart = len(ordered) == 1
        for previous, current in zip(ordered, ordered[1:]):
            gap = (
                separation(previous, current) if from_top else separation(current, previous)
            )
            if limit <= 0 or gap > limit:
                set_apart = True
                break
            if len(collected) >= _BOILERPLATE_EDGE_MAX_LINES:
                break
            collected.append(current)
        # Furniture is *set apart* from the body by white space. If the walk
        # reached its cap without ever finding that separation, this edge has no
        # distinct band — it is simply where the body text begins — and widening
        # here would offer real content to the boilerplate detector. A page
        # whose genuine section heading sits at the top edge is exactly this
        # case, so fall back to the outermost line alone.
        return collected if set_apart else [ordered[0]]

    seen: list[_Line] = []
    for line in band(from_top=True) + band(from_top=False):
        if not any(line is chosen for chosen in seen):
            seen.append(line)
    return seen


def _is_edge_line(line: _Line, lines: list[_Line]) -> bool:
    return any(candidate is line for candidate in _edge_lines(lines))


def _detect_boilerplate(page_lines: list[list[_Line]]) -> set[str]:
    """Running headers/footers occupy a fixed page-edge position on most pages.

    Restricting candidates to the page-edge band is what separates a true
    running header from a recurring in-body subheading, which drifts by a line
    or two depending on how long the preceding title is and is real content.
    """

    total_pages = len(page_lines)
    if total_pages < _BOILERPLATE_MIN_PAGES:
        return set()

    counts: Counter[str] = Counter()
    for lines in page_lines:
        seen: set[str] = set()
        for line in _edge_lines(lines):
            normalized = _normalize_line(line.text)
            if normalized and normalized not in seen:
                seen.add(normalized)
                counts[normalized] += 1

    threshold = max(_BOILERPLATE_MIN_PAGES, int(total_pages * _BOILERPLATE_MIN_PAGE_FRACTION))
    return {text for text, count in counts.items() if count >= threshold}


_SENTENCE_END = (".", "!", "?", "؟", "。")


def _classify_line(
    line: _Line, body_size: float, *, smallest_body_size: float | None = None
) -> ElementType:
    """Classify one line against the band of sizes the document sets text in.

    `body_size` is the top of that band and `smallest_body_size` the bottom.
    They differ only in a document with more than one text class; when the
    bottom is not given the band collapses to a point and this behaves exactly
    as a single modal size did.

    The two ends answer different questions, which is why one number cannot
    serve both. "Is this set larger than the body?" has to be asked against the
    largest text class, or the larger of two co-equal classes reads as a heading
    on every line. "Is this set smaller than the body?" has to be asked against
    the smallest, or the smaller of those two classes reads as furniture and is
    barred from carrying a heading at all.
    """

    stripped = line.text.strip()
    if not stripped:
        return "other"

    body_floor = body_size if smallest_body_size is None else smallest_body_size

    # Font size is strong, direct evidence of a structural heading, so it is
    # weighed before the weaker signals rather than after them.
    #
    # It has to outrank the list-marker test in particular. A document that
    # numbers its sections produces lines that carry a list marker and are also
    # set larger than the body, and deciding on the marker alone calls every one
    # of them a list item. The consequence is not a mislabelled element: a list
    # item is merged into the text that follows it, so the section title and its
    # first paragraph fuse into one, and the document loses every boundary
    # between its numbered sections at once. What separates an enumerated
    # heading from an enumerated list item is how it is set, not how it is
    # numbered.
    larger_font = body_size > 0 and line.size >= body_size + 0.6

    if _LIST_MARKER_RE.match(stripped):
        # Whether an enumerated line is a list item or a section heading cannot
        # be settled from the line alone. It is not decided by the shape of the
        # marker — a document is free to number its sections — and it cannot be
        # decided against the document's modal size either, because a document
        # with two co-equal text classes (a bilingual one, say) has a modal that
        # describes neither. It is settled in `_classify_lines`, which can see
        # what the line introduces.
        return "list_item"

    # Size is strong evidence of a heading, but it is not the only evidence, and
    # the weaker text-pattern rule below already refuses a line the document has
    # terminated as a sentence. Asking the same question here is not a second
    # rule; it is the same rule applied consistently. A heading names what
    # follows it and stops where its words stop, so a finished sentence is prose
    # however it is set — a bullet drawn a few points larger than its paragraph,
    # a pull quote, an emphasised warning. Left unasked, each of these becomes
    # the section that every record printed after it is filed under.
    if larger_font and len(stripped) <= 200 and not stripped.endswith(_SENTENCE_END):
        return "heading"

    # The text-pattern fallback ("ALL CAPS", "Article 12", "7.2.1 ...") is much
    # weaker, so it carries two guards.
    #
    # The first is size. A heading is never set smaller than the text it
    # introduces; type set below body size is furniture — a running footer, a
    # page number, a copyright line, a caption. Without this, a footer in small
    # capitals satisfies the ALL CAPS rule, becomes a heading, and is then
    # installed as the section label for everything printed after it on every
    # page it appears on. That reads as full section coverage while naming
    # something that is not a section at all.
    #
    # The second is punctuation: a line that ends in sentence punctuation is
    # content, not a label. Without it, a short numbered provision such as
    # "2.1 Such jobs shall not be potentially harmful to their health." is
    # classified as a heading, which both loses it as a candidate policy and
    # then mislabels the `section` of every element that follows it. The
    # asymmetry favours guessing "paragraph": a heading demoted to a paragraph
    # is merely a redundant element, whereas a rule promoted to a heading
    # disappears from extraction entirely.
    # `_modal_body_size` and the parser's own sizes both carry floating-point
    # noise, so a heading set at exactly body size can measure a hair under it.
    # The floor therefore uses the same margin that `larger_font` uses above:
    # only type that is *distinctly* smaller than the body is furniture. Without
    # the margin this rule demotes genuine section headings that happen to be
    # set at body size and distinguished by weight or spacing instead — which is
    # over-suppression, and costs real sections rather than false ones.
    smaller_font = body_floor > 0 and line.size < body_floor - 0.6
    if not smaller_font and len(stripped) <= 90 and not stripped.endswith(_SENTENCE_END):
        if stripped.isupper() or _PROVISION_HEADING_RE.match(stripped):
            return "heading"
    return "paragraph"


def _fills_the_measure(line: _Line, lines: list[_Line]) -> bool:
    """Whether a line runs as wide as the running text around it.

    A heading is set narrower than the body it introduces — it stops where its
    words stop, while a line of prose runs to the edge of the column and wraps.
    Comparing against the widest running text on the same page keeps the
    yardstick local, so a document with two columns or two text sizes is judged
    against its own measure rather than an assumed one.
    """

    measure = _column_measure(lines)
    if measure <= 0:
        return False
    return (line.x1 - line.x0) >= measure * _FILLS_THE_MEASURE_RATIO


def _is_uppercase_throughout(text: str) -> bool:
    """Whether every cased character is uppercase, and there is at least one.

    Case is the third way a document can set a line apart, alongside size and
    space, and it is the one that survives when the other two are unavailable —
    a section marker set at body size, tight against the line above it, is still
    marked as a heading by being the only capitalised thing on the page.

    Judged from Unicode character properties, so a script that does not
    distinguish case answers False and the decision falls to the other signals,
    rather than to an assumption about which alphabet is in use.
    """

    saw_a_cased_character = False
    for char in text:
        if char.islower():
            return False
        if char.isupper():
            saw_a_cased_character = True
    return saw_a_cased_character


def _starts_lowercase(text: str) -> bool:
    """Whether the first cased character is a lowercase one.

    Read from Unicode character properties rather than an alphabet, so it
    carries no assumption about which script it is looking at. A line opening
    with a bracket, quotation mark, digit or dash is judged on the first cased
    letter that follows, because the opening mark is not itself evidence either
    way.

    In a script with no case distinction the answer is always False, which is
    the honest one: such a script simply does not mark a continuation this way,
    and the decision falls to the other signals rather than to a guess.
    """

    for char in text.strip():
        if char.isupper():
            return False
        if char.islower():
            return True
    return False


def _classify_lines(
    lines: list[_Line], body_size: float, *, smallest_body_size: float | None = None
) -> list[ElementType]:
    """Classify a page's lines, then demote headings the next line contradicts.

    `_classify_line` sees one line at a time, and the font-size rule inside it
    is trusted on its own. That is right for a genuine label and wrong for a
    sentence that merely happens to be set in a larger face — and AD-103 has
    both. These two lines were classified `heading`:

        "The housing allowance per calendar year (12 months) is calculated as twice"
        "The housing allowance is limited to one employee of the married couple"

    Neither is a label. Each is the first half of a sentence whose second half
    is the following line ("the monthly basic salary up to a maximum of:",
    "(husband and wife). In the case of..."). Classified as headings they could
    never be merged — `_continues_previous` requires a paragraph or list item
    as the predecessor — so the sentence stayed cut, and the formulator
    reconstructed the governing sentence for the orphaned half and produced a
    duplicate rule.

    A following line that begins lowercase, or with an opening bracket, is
    direct evidence that the line before it did not finish. A heading does not
    have a continuation. So the pair overrules the font.

    This mirrors the guard the text-pattern branch of `_classify_line` already
    carries, and for the reason stated there: a heading demoted to a paragraph
    is a redundant element, whereas a rule promoted to a heading disappears
    from extraction entirely.
    """

    kinds = [
        _classify_line(line, body_size, smallest_body_size=smallest_body_size)
        for line in lines
    ]

    # Promote enumerated lines that the document sets above the content they
    # introduce. A document that numbers its sections produces lines carrying a
    # list marker that are nonetheless headings, and calling them list items
    # does not merely mislabel them: `_build_blocks` merges a list item with the
    # text that follows, so the section title and its opening paragraph fuse
    # into a single element and every boundary between numbered sections is
    # lost at once.
    #
    # The yardstick is local, not document-wide. A heading is set larger than
    # the text it introduces; a bullet in a list of bullets is set the same as
    # its siblings. Comparing against the neighbour rather than the document's
    # modal size is what keeps this honest in a document with two co-equal text
    # classes, where the modal describes one of them and slanders the other.
    for index, kind in enumerate(kinds):
        if kind != "list_item":
            continue
        stripped = lines[index].text.strip()
        if len(stripped) > 200 or stripped.endswith(_SENTENCE_END):
            continue
        following = next(
            (
                line
                for line in lines[index + 1 :]
                if line.text.strip() and not line.is_boilerplate
            ),
            None,
        )
        if following is not None and lines[index].size >= following.size + 0.6:
            kinds[index] = "heading"

    # Demote headings that the next line contradicts.
    #
    # The snapshot is taken first because the pass after this one needs to know
    # which lines were heading-like on their own typography, not which ones
    # survived this pass. A heading that wraps loses its first line here — the
    # second line begins lowercase, so the first is demoted — and without the
    # snapshot the wrapped pair is indistinguishable from a sentence tail.
    kinds_before_next_line_demotion = list(kinds)
    for index, kind in enumerate(kinds):
        if kind != "heading":
            continue
        nxt = next(
            (
                line.text.strip()
                for line in lines[index + 1 :]
                if line.text.strip() and not line.is_boilerplate
            ),
            "",
        )
        if not nxt:
            continue
        # A hard boundary in the next line means the heading stands: a new
        # provision or list item is its own start, not a continuation.
        if _PROVISION_START_RE.match(nxt) or _LIST_MARKER_RE.match(nxt):
            continue
        first = nxt[:1]
        if first.islower() or first in "([":
            kinds[index] = "paragraph"

    # Demote headings that are really the tail of the line before them.
    #
    # This runs after the forward-looking pass above, not before it, because it
    # reads the settled type of the preceding line. Run first, it would see
    # lines that pass is about to demote and mistake a run of misclassified
    # paragraph lines for a heading that wrapped.
    #
    # The pass above asks whether the *next* line contradicts a heading. Nothing
    # asked whether the heading contradicts the line *before* it, and that
    # asymmetry had no justification: a heading begins something, so it is never
    # itself a continuation. Where the classifier types an orphaned tail as a
    # heading the damage compounds, because `_build_blocks` breaks a block at a
    # heading and `_continues_previous` will not merge across one — so the
    # sentence is cut, the halves can never be rejoined, and the tail carries no
    # section, since a heading is a section rather than being in one. A clause
    # then arrives at extraction without the words that gave its references an
    # antecedent.
    #
    # The test is whether the document sets the line apart, which it can do
    # three ways: by size, by space, or by case. A genuine heading is set larger
    # than what precedes it, separated from it by more than the running leading,
    # or cased distinctly from the running text — often more than one at once. A
    # line that continues the sentence above it is set apart by none of them.
    # All three are read from the page's own geometry and from Unicode
    # character properties, so nothing here depends on vocabulary, script or
    # layout — and in a script that marks neither case nor a sentence end this
    # way, the rule simply declines to fire rather than guessing.
    leading = _typical_leading(lines)
    for index, kind in enumerate(kinds):
        if kind != "heading":
            continue
        previous_index = next(
            (
                candidate
                for candidate in range(index - 1, -1, -1)
                if lines[candidate].text.strip() and not lines[candidate].is_boilerplate
            ),
            None,
        )
        if previous_index is None:
            continue
        previous_line = lines[previous_index]
        # A heading that runs onto a second line is a heading that wrapped, not
        # an orphan: the line it continues did the setting apart on its behalf,
        # so the second line carries none of the three signals itself. Repairing
        # that is not this rule's job — demoting it merges a real heading into
        # the body and a section disappears.
        #
        # Two things must hold for a line to be read that way, and both are
        # relational rather than absolute. It must not close a sentence, because
        # a heading does not end in a full stop. And it must not fill the
        # measure: a heading is set narrower than the running text, so a line
        # that runs the full column width is prose regardless of what precedes
        # it. Width is used rather than character count because a character
        # count means different things in different scripts, while the width of
        # the text block does not.
        wrapped_from_a_heading = (
            kinds_before_next_line_demotion[previous_index] == "heading"
            and _ends_open(lines[index].text)
            and not _fills_the_measure(lines[index], lines)
        )
        if wrapped_from_a_heading:
            continue
        looks_like_a_continuation = _starts_lowercase(lines[index].text) or _ends_open(
            previous_line.text
        )
        if not looks_like_a_continuation:
            continue
        set_apart_by_size = lines[index].size >= previous_line.size + 0.6
        gap = lines[index].top - previous_line.bottom
        set_apart_by_space = leading > 0 and gap > leading * 1.6
        set_apart_by_case = _is_uppercase_throughout(
            lines[index].text
        ) and not _is_uppercase_throughout(previous_line.text)
        if set_apart_by_size or set_apart_by_space or set_apart_by_case:
            continue
        kinds[index] = "paragraph"

    return kinds


def _build_blocks(
    lines: list[_Line],
    table_blocks: list[_Block],
    body_size: float,
    *,
    smallest_body_size: float | None = None,
) -> list[_Block]:
    """Group a page's lines into logical blocks, then interleave table rows.

    Blocks break on: a heading, the start of a new list item, entering or
    leaving a table region, and a vertical gap larger than normal leading
    (which is how PDFs express a paragraph break). They deliberately do *not*
    break on line count — that was the old fixed-four-line behaviour the spec
    prohibits (section 12).
    """

    blocks: list[_Block] = []
    current: _Block | None = None
    previous: _Line | None = None
    leading = _typical_leading(lines)
    # Classified for the whole page up front, so a heading can be demoted by the
    # line that follows it — see `_classify_lines`.
    kinds = _classify_lines(lines, body_size, smallest_body_size=smallest_body_size)

    for index, line in enumerate(lines):
        if line.is_boilerplate or not line.text.strip():
            previous = line
            continue
        if line.in_table:
            # Table content is emitted from the structured table extraction
            # instead, so it is not also emitted as loose paragraph text.
            current = None
            previous = line
            continue

        kind = kinds[index]
        gap = (line.top - previous.bottom) if previous is not None else 0.0
        paragraph_break = previous is not None and leading > 0 and gap > leading * 1.6

        starts_new_block = (
            current is None
            or kind == "heading"
            or current.element_type == "heading"
            or kind == "list_item"
            or paragraph_break
        )
        if starts_new_block:
            current = _Block(element_type=kind)
            blocks.append(current)
        current.lines.append(line)
        previous = line

    blocks.extend(table_blocks)
    blocks.sort(key=lambda block: (block.top, block.page))
    return [block for block in blocks if block.lines]


def _typical_leading(lines: list[_Line]) -> float:
    """Median vertical gap between consecutive lines, used to spot paragraph breaks."""

    gaps: list[float] = []
    for previous, current in zip(lines, lines[1:]):
        gap = current.top - previous.bottom
        if 0 <= gap < 60:
            gaps.append(gap)
    if not gaps:
        return 0.0
    gaps.sort()
    return gaps[len(gaps) // 2]


def _join_lines(lines: list[_Line]) -> tuple[str, list[Transformation]]:
    """Concatenate line texts, recording every transformation applied.

    Two joins exist, both deterministic: a single space between lines, and a
    *hyphen-preserving* join when a line ends on a hyphen ("employ-" / "ment"
    becomes "employ-ment", not "employment").

    WHY THE HYPHEN IS KEPT
    ----------------------
    Spec section 30 permits dehyphenation provided it is deterministic, coded
    and recorded. It is not implemented, because it cannot be done correctly:
    a line ending in "non-" followed by "renewal" is indistinguishable from
    "employ-" followed by "ment", yet the first is a real compound hyphen and
    the second is visual line-break hyphenation. Both only ever occur at a
    right margin, so no geometric signal separates them, and most PDFs encode
    both as U+002D rather than using a soft hyphen.

    Removing the hyphen therefore risks producing "nonrenewal" — a token that
    does not appear anywhere in the document — which is precisely what
    INVARIANT 6 exists to prevent. Keeping it produces "employ-ment", which is
    merely ugly: every character still came from the source, a reader and a
    model both parse it correctly, and nothing is fabricated. Given that
    asymmetry, preservation is the only defensible default.

    If dehyphenated text is ever needed for search recall, spec section 30's
    "reversible or source-mapped" clause allows deriving a *separate* field
    from this one. It must not mutate canonical text, which section 28 requires
    to stay stable because every recorded offset indexes into it.
    """

    parts: list[str] = []
    transformations: list[Transformation] = []
    for index, line in enumerate(lines):
        text = line.text
        if index == 0:
            parts.append(text)
            continue
        previous = parts[-1]
        if (
            previous.endswith("-")
            and len(previous) > 1
            and previous[-2].isalpha()
            and text[:1].islower()
        ):
            parts.append(text)
            if "line_break_hyphen_join" not in transformations:
                transformations.append("line_break_hyphen_join")
            continue
        parts.append(" ")
        parts.append(text)
        if "line_join_space" not in transformations:
            transformations.append("line_join_space")
    return "".join(parts), transformations


def _ends_open(text: str) -> bool:
    stripped = text.rstrip()
    if not stripped:
        return False
    return not stripped.endswith(_TERMINAL_PUNCTUATION)


def _continues_previous(previous: _Block, nxt: _Block, *, same_page: bool = False) -> bool:
    """Decide whether a block continues the previous one.

    Uses several structural signals together rather than one heuristic (spec
    section 9), because any single signal misfires: plenty of legitimate
    paragraphs start lowercase, and plenty of continuations start with a capital
    proper noun.

    `same_page` tightens the last rule. Across a page break, an open-ended
    previous block followed by anything that is not a heading or a provision is
    treated as a continuation, because a sentence does not normally end at a
    page boundary without punctuation. Within a page that reasoning does not
    hold: blocks are also split for layout reasons — spacing, font change, a new
    column — and a capitalised block after an unpunctuated one is commonly a
    genuine new paragraph. So within a page the next block must *itself* look
    like a continuation: lowercase, or opening with punctuation such as a
    bracket.

    This was originally reachable only for page-leading blocks, which left
    mid-page breaks unmerged. Two clauses in AD-103 begin mid-sentence for that
    reason:

        "the monthly basic salary up to a maximum of:"
        "(husband and wife). In the case of a married couple are employed by FBSU"

    the second being the tail of "The housing allowance is limited to one
    employee of the married couple (husband and wife). In the case of…". The
    formulator then reconstructs the governing sentence from inherited context
    for the orphaned half and produces a rule the preceding clause already
    produced — two mid-sentence cuts, two duplicate rule pairs, exact
    correspondence.
    """

    if previous.element_type not in ("paragraph", "list_item"):
        return False
    if nxt.element_type not in ("paragraph",):
        return False
    if previous.table_id or nxt.table_id:
        return False

    previous_text = " ".join(line.text for line in previous.lines).strip()
    next_text = " ".join(line.text for line in nxt.lines).strip()
    if not previous_text or not next_text:
        return False

    # Hard boundary signals win outright.
    if _PROVISION_START_RE.match(next_text) or _LIST_MARKER_RE.match(next_text):
        return False

    if not _ends_open(previous_text):
        return False

    first_char = next_text[:1]
    if first_char.islower() or not first_char.isalpha():
        return True
    # An open-ended previous paragraph plus a non-heading, non-provision
    # continuation is still a continuation even when capitalized, since the
    # break mid-sentence is itself strong evidence — but only across a page
    # boundary, where a break carries no layout meaning.
    return not same_page


def _assemble_elements(page_blocks: list[list[_Block]]) -> list[CanonicalElement]:
    """Flatten pages into one ordered element list, merging cross-page blocks."""

    merged: list[_Block] = []
    for blocks in page_blocks:
        for index, block in enumerate(blocks):
            # `index == 0` used to gate this entirely, so only a page-leading
            # block could ever be merged and a sentence cut mid-page stayed cut.
            # Both signals now run, with the within-page case held to the
            # stricter test — see `_continues_previous`.
            if merged and _continues_previous(merged[-1], block, same_page=index > 0):
                merged[-1].lines.extend(block.lines)
                continue
            merged.append(block)

    elements: list[CanonicalElement] = []
    section: str | None = None
    for order, block in enumerate(merged, start=1):
        if block.cell_text is not None:
            text = block.cell_text
            transformations: list[Transformation] = ["table_cell_join"]
        else:
            text, transformations = _join_lines(block.lines)
        text = text.strip()
        if not text:
            continue
        if len({line.page for line in block.lines}) > 1:
            transformations.append("cross_page_join")

        if block.element_type == "heading":
            section = text

        elements.append(
            CanonicalElement(
                element_id=f"E{len(elements) + 1:06d}",
                element_type=block.element_type,
                logical_order=len(elements),
                text=text,
                section=None if block.element_type == "heading" else section,
                source_fragments=_fragments_for(block.lines),
                transformations=transformations,
                table_id=block.table_id,
                table_headers=block.table_headers,
            )
        )
    return elements


def _fragments_for(lines: list[_Line]) -> list[SourceFragment]:
    """One fragment per contiguous run of lines on the same page.

    Adjacent lines on one page collapse into a single span so a cross-page
    element reports exactly one fragment per page it touches, which is the
    shape spec section 25 specifies and section 27 validates against.
    """

    fragments: list[SourceFragment] = []
    for line in lines:
        if fragments and fragments[-1].page == line.page and fragments[-1].end_offset <= line.start_offset:
            previous = fragments[-1]
            fragments[-1] = SourceFragment(
                page=previous.page,
                start_offset=previous.start_offset,
                end_offset=line.end_offset,
                text=previous.text + "\n" + line.text
                if line.start_offset == previous.end_offset + 1
                else previous.text + " " + line.text,
            )
        else:
            fragments.append(line.fragment)
    return fragments

# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def ingest_docx(storage_path: str | Path, document_id: str = "") -> CanonicalDocument:
    """Parse a DOCX into the same canonical representation as a PDF.

    DOCX has no page structure — pagination is decided by the renderer, not
    stored in the file — so the whole document is modelled as one logical page
    whose ``raw_text`` is every block joined in document order. Offsets index
    into that string, so provenance stays exactly as verifiable as it is for a
    PDF, and downstream consumers need no special case.

    The previous DOCX path turned each table row into prose of the form
    ``"Tier: 2; Limit: 5000"``. That text appears nowhere in the document, so
    any passage "quoted" from it was fabricated by definition — the very thing
    INVARIANT 6 exists to prevent. Rows now keep their cell values verbatim and
    carry their headers as structured data instead.
    """

    path = Path(storage_path)
    try:
        document = DocxDocument(str(path))
    except Exception as exc:
        raise IngestionError(f"cannot open DOCX {path.name}: {exc}") from exc

    diagnostics: list[IngestionDiagnostic] = []
    raw_parts: list[str] = []
    offset = 0
    elements: list[CanonicalElement] = []
    section: str | None = None
    table_index = 0

    def _emit(
        text: str,
        element_type: ElementType,
        table_id: str | None = None,
        headers: list[str] | None = None,
        transformations: list[Transformation] | None = None,
        raw_line: str | None = None,
    ) -> None:
        nonlocal offset, section
        line = raw_line if raw_line is not None else text
        if not line.strip():
            return
        start = offset
        end = start + len(line)
        raw_parts.append(line)
        offset = end + 1  # the newline joining blocks

        if element_type == "heading":
            section = text
        elements.append(
            CanonicalElement(
                element_id=f"E{len(elements) + 1:06d}",
                element_type=element_type,
                logical_order=len(elements),
                text=text,
                section=None if element_type == "heading" else section,
                source_fragments=[
                    SourceFragment(page=1, start_offset=start, end_offset=end, text=line)
                ],
                transformations=transformations or [],
                table_id=table_id,
                table_headers=headers,
            )
        )

    for block in _iter_docx_blocks(document):
        if isinstance(block, DocxParagraph):
            text = block.text.strip()
            if not text:
                continue
            _emit(text, _classify_docx_paragraph(block, text))
        elif isinstance(block, DocxTable):
            table_index += 1
            table_id = f"t{table_index}"
            try:
                rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
            except Exception as exc:
                diagnostics.append(
                    IngestionDiagnostic(
                        code="table_parse_failed",
                        severity="warning",
                        detail=f"table {table_id}: {exc}",
                    )
                )
                continue
            if not rows:
                continue
            genuine = _is_genuine_table(rows)
            # Evidenced, not assumed. This path used to read row 0 as the header
            # whenever any of its cells held text and then drop it with
            # `rows[1:]` -- the exact defect fa27428 removed from the PDF path,
            # left live here, deleting a row of every DOCX table on no evidence
            # and telling nobody. The verdict and its reason now come from the
            # one place both parsers ask.
            has_headers = False
            if genuine:
                has_headers, header_diagnostic = _column_labels_for(
                    rows, table_id=table_id, page=1
                )
                if header_diagnostic is not None:
                    diagnostics.append(header_diagnostic)
            # `None` rather than `[]`: no row stated column labels, which is not
            # the same as a header row that was blank.
            headers = rows[0] if has_headers else None
            data_rows = rows[1:] if has_headers else rows
            for row in data_rows:
                if not any(row):
                    continue
                if genuine:
                    _emit(
                        " | ".join(row),
                        "table_row",
                        table_id=table_id,
                        headers=headers,
                        transformations=["table_cell_join"],
                    )
                else:
                    # A single-column "table" is a layout device, so its cells
                    # are ordinary paragraphs.
                    for cell in row:
                        if cell.strip():
                            _emit(cell, "paragraph")

    canonical = CanonicalDocument(
        document_id=document_id or path.stem,
        page_count=1,
        pages=[CanonicalPage(page=1, raw_text="\n".join(raw_parts))],
        elements=elements,
        parser=DOCX_PARSER_NAME,
        diagnostics=diagnostics,
    )
    _append_document_diagnostics(canonical)
    return canonical


def _classify_docx_paragraph(block: DocxParagraph, text: str) -> ElementType:
    """Prefer the document's own declared style over guessing from the text."""

    try:
        style_name = (block.style.name if block.style else "") or ""
    except Exception:
        style_name = ""
    lowered = style_name.lower()
    if lowered.startswith("heading") or lowered in ("title", "subtitle"):
        return "heading"
    if "list" in lowered:
        return "list_item"
    if _LIST_MARKER_RE.match(text):
        return "list_item"
    if len(text) <= 90 and not text.endswith(_SENTENCE_END):
        if text.isupper() or _PROVISION_HEADING_RE.match(text):
            return "heading"
    return "paragraph"


def _iter_docx_blocks(document):
    """Yield paragraphs and tables in true document order.

    python-docx exposes ``paragraphs`` and ``tables`` as separate collections
    with no interleaving, so walking the XML body is the only way to preserve
    the order a reader sees — and order is what makes a table row attributable
    to the section heading above it.
    """

    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield DocxParagraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield DocxTable(child, document)
