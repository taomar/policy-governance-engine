"""Layout-aware document ingestion into a canonical representation.

Implements the PDF ingestion architecture spec (docs/specs/pdf-ingestion-architecture-v1.md).

WHY THIS EXISTS
---------------
The previous ingestion (``document_extraction._extract_pdf``) iterated page by
page and emitted paragraphs scoped to a single page, falling back to grouping
lines in fixed batches of four when a page had no blank-line breaks. Two
architectural consequences followed, and both are visible in extracted output:

1. A paragraph that crossed a page boundary was split into two clauses, so a
   condition could be separated from its consequence, or a rule from its
   exception. The extractor then saw half a rule and formalized half a rule.
2. Fixed-size line grouping made *chunk boundaries into policy boundaries*,
   which the spec forbids outright (sections 12, 13).

Neither was a bug in one document — per spec section 58, the same failure was
possible in every document, so the fix belongs at the ingestion layer rather
than in a per-document workaround.

HOW IT WORKS
------------
Text is reconstructed from ``page.extract_words()`` rather than
``page.extract_text()``. That is deliberate: words carry geometry (top, x0,
font size), which makes several things decidable that are otherwise guesswork —
which lines belong to a table (bbox containment), which lines are headings
(font size relative to the document's modal body size), where column gutters
fall, and where line boundaries actually are. Because this module *constructs*
the canonical page string itself, every offset it records is exact by
construction rather than recovered by searching, satisfying INVARIANT 4 and
INVARIANT 5.

ARBITRARY UPLOADS
-----------------
Files arrive from users, so nothing here may assume well-formed input. The
degenerate cases that actually occur — scans with no text layer, encrypted
files, two-column layouts, rotated watermark text, right-to-left script,
bordered callouts that ruling-line detection mistakes for tables — are each
handled explicitly and, where they cannot be handled, reported as an
``IngestionDiagnostic`` rather than silently reducing coverage (INVARIANT 9).
A document that yields little text because it is a scan must never be
indistinguishable from a document that genuinely contains little policy.

WHAT IT MUST NEVER DO
---------------------
Insert a word. Cross-page reconstruction concatenates existing fragments in
source order (spec section 10). The only text transformations permitted are
deterministic, code-performed, and recorded on the element:
``line_join_space``, ``line_break_hyphen_join``, ``cross_page_join``,
``table_cell_join``.
"""
from __future__ import annotations

import re
import unicodedata
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from policy_platform.contracts.canonical_document import (
    CanonicalDocument,
    CanonicalElement,
    CanonicalPage,
    ElementType,
    IngestionDiagnostic,
    SourceFragment,
    Transformation,
)

PARSER_NAME = "pdfplumber-layout-v1"
DOCX_PARSER_NAME = "python-docx-layout-v1"

# Words on the same visual line rarely differ in `top` by more than a couple of
# points; anything larger is a genuinely new line. Kept tight because legal PDFs
# often set dense leading, where a loose tolerance merges consecutive lines.
_Y_TOLERANCE = 3.0

# Matches pdfplumber's own default word-splitting behaviour poorly on these
# documents: the default merges adjacent words that lack a real space, yielding
# artifacts like "thispolicy". 1.0 recovers the gap.
_X_TOLERANCE = 1.0

_TERMINAL_PUNCTUATION = ('.', '!', '?', ':', ';', '"', '”', '’', "'", ')', '،', '؛', '.')

_LIST_MARKER_RE = re.compile(
    r"^\s*(?:"
    r"\(\d{1,3}\)"          # (1)
    r"|\d{1,3}[.)]"          # 1.  1)
    r"|\(?[a-zA-Z][.)]"      # a.  (a)
    r"|[ivxIVX]{1,5}[.)]"    # iv.
    r"|[\u2022\u25AA\u25E6\u00B7\u2023\u2043*\-\u2013\u2014]"  # bullets/dashes
    r")\s+"
)

# A new numbered provision is a hard boundary signal: it must never be merged
# onto the tail of the previous page's paragraph (spec section 9). Deliberately
# loose — for boundary detection, over-matching is safe (it only prevents a
# join), whereas under-matching silently welds two provisions together.
_PROVISION_START_RE = re.compile(
    r"^\s*(?:"
    r"(?:Article|ARTICLE|Section|SECTION|Chapter|CHAPTER|Clause|CLAUSE|Part|PART|Annex|ANNEX|Appendix|APPENDIX)\b"
    r"|\u0627\u0644\u0645\u0627\u062F\u0629\b"   # Arabic "the Article"
    r"|\u0627\u0644\u0641\u0635\u0644\b"          # Arabic "the Chapter"
    r"|\u0627\u0644\u0628\u0627\u0628\b"          # Arabic "the Part"
    r"|\d{1,3}(?:\.\d{1,3}){1,4}\s"               # 7.2.1
    r")"
)

# Whether a line is a provision *heading* is a much stricter question than
# whether it *starts* one, and needs its own pattern. "Article (81)" is a
# heading; "Article (81) of the Labor Law during the training term or within"
# is the middle of a sentence that merely happens to begin with a citation.
# Distinguishing them matters twice over: the second form is a real policy
# clause that would vanish from extraction, and it would also be installed as
# the `section` label for every element after it.
_PROVISION_HEADING_RE = re.compile(
    r"^\s*(?:"
    r"Article|ARTICLE|Section|SECTION|Chapter|CHAPTER|Clause|CLAUSE|Part|PART|Annex|ANNEX|Appendix|APPENDIX"
    r"|\u0627\u0644\u0645\u0627\u062F\u0629|\u0627\u0644\u0641\u0635\u0644|\u0627\u0644\u0628\u0627\u0628"
    r")\s*"
    r"\(?\s*[\dIVXLC]{1,6}\s*\)?"                      # 12, (12), IV
    r"(?:\s*(?:bis|BIS|ter|TER|\u0645\u0643\u0631\u0631))?"   # "11 bis"
    r"\s*(?:[:\-\u2013\u2014.]\s*\S.{0,140})?"          # optional ": Definitions"
    r"\s*$"
)

_BOILERPLATE_MIN_PAGE_FRACTION = 0.3
_BOILERPLATE_MIN_PAGES = 3
_BOILERPLATE_EDGE_WINDOW = 1
_DIGIT_RE = re.compile(r"\d+")


@dataclass
class _Line:
    """One visual line of text, with the geometry needed to classify it."""

    text: str
    top: float
    bottom: float
    x0: float
    x1: float
    size: float
    page: int
    start_offset: int = 0
    end_offset: int = 0
    in_table: str | None = None  # table_id when the line sits inside a table bbox
    is_boilerplate: bool = False

    @property
    def fragment(self) -> SourceFragment:
        return SourceFragment(
            page=self.page,
            start_offset=self.start_offset,
            end_offset=self.end_offset,
            text=self.text,
        )


@dataclass
class _Block:
    """A run of lines that form one logical unit *within a single page*.

    Cross-page merging happens afterwards, on blocks, so that the decision to
    join has both candidates fully assembled rather than guessing line by line.
    """

    element_type: ElementType
    lines: list[_Line] = field(default_factory=list)
    table_id: str | None = None
    table_headers: list[str] | None = None
    cell_text: str | None = None  # set for table rows, whose text is not line-derived

    @property
    def top(self) -> float:
        return min(line.top for line in self.lines) if self.lines else 0.0

    @property
    def page(self) -> int:
        return self.lines[0].page if self.lines else 0


class IngestionError(RuntimeError):
    """Raised when a file cannot be ingested at all (encrypted, corrupt, unsupported)."""


def ingest_document(
    storage_path: str | Path, mime_type: str = "", document_id: str = ""
) -> CanonicalDocument:
    """Ingest any supported upload into the canonical representation."""

    path = Path(storage_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf" or "pdf" in mime_type:
        return ingest_pdf(path, document_id)
    if suffix in (".docx", ".doc") or "wordprocessingml" in mime_type:
        return ingest_docx(path, document_id)
    raise IngestionError(f"unsupported document type: {mime_type or suffix or path.name}")


def ingest_pdf(storage_path: str | Path, document_id: str = "") -> CanonicalDocument:
    """Parse a PDF into the canonical representation.

    Deterministic: the same bytes always produce the same element ids, order,
    and offsets (spec section 45). No model is involved at any point.
    """

    path = Path(storage_path)
    diagnostics: list[IngestionDiagnostic] = []
    pages: list[CanonicalPage] = []
    page_blocks: list[list[_Block]] = []
    all_sizes: Counter[float] = Counter()
    raw_pages: list[tuple[int, list[_Line], list[_Block]]] = []

    try:
        pdf = pdfplumber.open(path)
    except Exception as exc:  # encrypted, truncated, not actually a PDF
        raise IngestionError(f"cannot open PDF {path.name}: {exc}") from exc

    with pdf:
        if not pdf.pages:
            raise IngestionError(f"PDF {path.name} contains no pages")
        for page_index, page in enumerate(pdf.pages, start=1):
            try:
                lines, tables, page_diagnostics = _read_page(page, page_index)
            except Exception as exc:  # one bad page must not lose the document
                diagnostics.append(
                    IngestionDiagnostic(
                        code="page_parse_failed",
                        severity="error",
                        page=page_index,
                        detail=str(exc),
                    )
                )
                lines, tables, page_diagnostics = [], [], []
            diagnostics.extend(page_diagnostics)
            for line in lines:
                all_sizes[round(line.size, 1)] += 1
            raw_pages.append((page_index, lines, tables))

    body_size = _modal_body_size(all_sizes)
    boilerplate = _detect_boilerplate([lines for _, lines, _ in raw_pages])

    for page_index, lines, table_blocks in raw_pages:
        removed: list[str] = []
        for line in lines:
            if _normalize_line(line.text) in boilerplate and _is_edge_line(line, lines):
                line.is_boilerplate = True
                removed.append(line.text)
        # INVARIANT: raw_text keeps every character the parser produced,
        # including headers and footers. They are removed from the *logical*
        # flow only (spec section 11) — dropping them from raw_text would
        # invalidate every offset already recorded against it.
        raw_text = "\n".join(line.text for line in lines)
        pages.append(CanonicalPage(page=page_index, raw_text=raw_text, removed_boilerplate=removed))
        page_blocks.append(_build_blocks(lines, table_blocks, body_size))

    elements = _assemble_elements(page_blocks)

    document = CanonicalDocument(
        document_id=document_id or path.stem,
        page_count=len(pages),
        pages=pages,
        elements=elements,
        parser=PARSER_NAME,
        diagnostics=diagnostics,
    )
    _append_document_diagnostics(document)
    return document


def _append_document_diagnostics(document: CanonicalDocument) -> None:
    """Report degradations that would otherwise look like a sparse document."""

    empty_pages = [page.page for page in document.pages if not page.raw_text.strip()]
    if empty_pages and len(empty_pages) >= max(1, len(document.pages) // 2):
        document.diagnostics.append(
            IngestionDiagnostic(
                code="no_text_layer",
                severity="error",
                detail=(
                    f"{len(empty_pages)} of {len(document.pages)} pages contain no extractable "
                    "text. The file is most likely a scan and needs OCR before it can be used "
                    "for policy extraction."
                ),
            )
        )
    elif empty_pages:
        document.diagnostics.append(
            IngestionDiagnostic(
                code="empty_pages",
                severity="warning",
                detail=f"pages with no extractable text: {empty_pages[:20]}",
            )
        )

    if not document.elements:
        document.diagnostics.append(
            IngestionDiagnostic(
                code="no_elements",
                severity="error",
                detail="ingestion produced no logical elements",
            )
        )
        return

    raw_chars = sum(
        len(_strip_whitespace(page.raw_text)) for page in document.pages
    )
    element_chars = sum(len(_strip_whitespace(element.text)) for element in document.elements)
    if raw_chars:
        ratio = element_chars / raw_chars
        if ratio < 0.85:
            document.diagnostics.append(
                IngestionDiagnostic(
                    code="low_coverage",
                    severity="warning",
                    detail=(
                        f"only {ratio:.1%} of source characters reached a logical element; "
                        "content may be trapped in unparsed regions"
                    ),
                )
            )

    failures = document.verify_fragments()
    if failures:
        document.diagnostics.append(
            IngestionDiagnostic(
                code="fragment_offsets_unresolvable",
                severity="error",
                detail=f"{len(failures)} source fragments do not resolve; first: {failures[0]}",
            )
        )

    rtl_chars = sum(
        1
        for element in document.elements
        for char in element.text
        if _is_rtl(char)
    )
    total_letters = sum(
        1 for element in document.elements for char in element.text if char.isalpha()
    )
    if total_letters and rtl_chars / total_letters > 0.2:
        document.diagnostics.append(
            IngestionDiagnostic(
                code="rtl_script_detected",
                severity="warning",
                detail=(
                    "right-to-left script detected. Words are ordered by horizontal position, "
                    "which may not match logical reading order for this script. No reordering "
                    "was applied because it cannot be verified against this source, and a wrong "
                    "reordering would corrupt text that is currently merely awkward."
                ),
            )
        )


def _strip_whitespace(text: str) -> str:
    return "".join(text.split())


def _is_rtl(char: str) -> bool:
    return unicodedata.bidirectional(char) in ("R", "AL")


def _read_page(page, page_index: int) -> tuple[list[_Line], list[_Block], list[IngestionDiagnostic]]:
    """Reconstruct one page's lines (with exact offsets) and its tables.

    Lines are built and offset *before* tables are resolved, because a table
    row's provenance is the span of real source lines it covers. Deriving row
    offsets any other way would mean inventing positions, and a source position
    that does not resolve is worse than no position at all — it makes an
    unverifiable extraction look verified.
    """

    diagnostics: list[IngestionDiagnostic] = []
    try:
        words = page.extract_words(x_tolerance=_X_TOLERANCE, extra_attrs=["size"])
    except Exception:  # malformed font metrics; retry without the extra attribute
        words = page.extract_words(x_tolerance=_X_TOLERANCE)

    # Rotated text is almost always a watermark, sidebar label, or stamp. Left
    # in the flow it interleaves with body text at arbitrary positions and
    # corrupts reading order, so it is excluded and reported rather than
    # silently mixed in.
    upright_words = [word for word in words if word.get("upright", True)]
    rotated_count = len(words) - len(upright_words)
    if rotated_count:
        diagnostics.append(
            IngestionDiagnostic(
                code="rotated_text_excluded",
                severity="info",
                page=page_index,
                detail=f"{rotated_count} rotated words excluded from the logical flow",
            )
        )

    if not upright_words:
        return [], [], diagnostics

    lines = _group_words_into_lines(upright_words, page_index)

    columns = _detect_columns(lines, float(getattr(page, "width", 0.0) or 0.0))
    if columns:
        diagnostics.append(
            IngestionDiagnostic(
                code="multi_column_layout",
                severity="info",
                page=page_index,
                detail=f"{len(columns)} columns detected; lines ordered per column",
            )
        )
        lines = _order_by_column(lines, columns)

    offset = 0
    for index, line in enumerate(lines):
        line.start_offset = offset
        line.end_offset = offset + len(line.text)
        offset = line.end_offset + (1 if index < len(lines) - 1 else 0)

    table_regions: list[tuple[str, tuple[float, float, float, float]]] = []
    table_blocks: list[_Block] = []
    try:
        found = page.find_tables()
    except Exception:  # pdfplumber can throw on degenerate ruling geometry
        found = []
    for table_index, table in enumerate(found, start=1):
        table_id = f"p{page_index}-t{table_index}"
        blocks = _table_to_blocks(table, table_id, page_index, lines)
        if not blocks:
            # Not a genuine table (see _is_genuine_table). Leaving the region
            # unregistered is the whole point: the lines then flow through
            # normal paragraph grouping and get reassembled correctly.
            continue
        table_regions.append((table_id, table.bbox))
        table_blocks.extend(blocks)

    for line in lines:
        line.in_table = _containing_table(line, table_regions)

    return lines, table_blocks, diagnostics


# A gutter narrower than this is word spacing, not a column separator.
_MIN_GUTTER_POINTS = 16.0
# Fraction of lines that may straddle a candidate gutter before it is rejected.
_MAX_GUTTER_CROSSING = 0.08
# Each column must carry at least this share of the page's lines.
_MIN_COLUMN_SHARE = 0.15


def _detect_columns(
    lines: list[_Line], page_width: float
) -> list[tuple[float, float]]:
    """Find vertical gutters that separate genuine text columns.

    Two-column layouts are common in regulations and handbooks, and grouping
    words by vertical position alone welds the left and right column of the
    same visual line into one string — silently producing text that reads as
    nonsense while looking perfectly well-formed. That failure is invisible
    downstream, so it has to be caught here.

    A candidate gutter is an interior band of x with no text in it. It is only
    accepted when almost no line crosses it (a real gutter is respected by
    every line) and both sides carry a meaningful share of the page's lines
    (which rejects the ragged right margin of ordinary prose).
    """

    if page_width <= 0 or len(lines) < 8:
        return []

    width = int(page_width) + 1
    occupied = bytearray(width)
    for line in lines:
        start = max(0, int(line.x0))
        end = min(width - 1, int(line.x1))
        for x in range(start, end + 1):
            occupied[x] = 1

    gaps: list[tuple[int, int]] = []
    run_start: int | None = None
    for x in range(width):
        if not occupied[x]:
            if run_start is None:
                run_start = x
        elif run_start is not None:
            gaps.append((run_start, x))
            run_start = None

    total = len(lines)
    accepted: list[tuple[int, int]] = []
    for start, end in gaps:
        if end - start < _MIN_GUTTER_POINTS:
            continue
        if start < page_width * 0.15 or end > page_width * 0.85:
            continue  # page margins, not a gutter
        centre = (start + end) / 2
        crossing = sum(1 for line in lines if line.x0 < centre < line.x1)
        if crossing > total * _MAX_GUTTER_CROSSING:
            continue
        left = sum(1 for line in lines if line.x1 <= centre)
        right = sum(1 for line in lines if line.x0 >= centre)
        if left < total * _MIN_COLUMN_SHARE or right < total * _MIN_COLUMN_SHARE:
            continue
        accepted.append((start, end))

    if not accepted:
        return []

    boundaries = [0.0]
    for start, end in accepted:
        boundaries.append((start + end) / 2)
    boundaries.append(page_width)
    return [(boundaries[i], boundaries[i + 1]) for i in range(len(boundaries) - 1)]


def _order_by_column(
    lines: list[_Line], columns: list[tuple[float, float]]
) -> list[_Line]:
    """Read each column top-to-bottom before moving to the next."""

    buckets: list[list[_Line]] = [[] for _ in columns]
    for line in lines:
        centre = (line.x0 + line.x1) / 2
        index = 0
        for candidate, (start, end) in enumerate(columns):
            if start <= centre < end:
                index = candidate
                break
        else:
            index = len(columns) - 1
        buckets[index].append(line)
    ordered: list[_Line] = []
    for bucket in buckets:
        ordered.extend(sorted(bucket, key=lambda item: (item.top, item.x0)))
    return ordered


def _group_words_into_lines(words: list[dict], page_index: int) -> list[_Line]:
    if not words:
        return []
    ordered = sorted(words, key=lambda w: (round(w["top"], 1), w["x0"]))
    lines: list[_Line] = []
    current: list[dict] = []
    current_top: float | None = None

    for word in ordered:
        if current_top is None or abs(word["top"] - current_top) <= _Y_TOLERANCE:
            if current_top is None:
                current_top = word["top"]
            current.append(word)
        else:
            lines.append(_line_from_words(current, page_index))
            current = [word]
            current_top = word["top"]
    if current:
        lines.append(_line_from_words(current, page_index))
    return lines


def _line_from_words(words: list[dict], page_index: int) -> _Line:
    ordered = sorted(words, key=lambda w: w["x0"])
    sizes = [float(w.get("size", 0.0) or 0.0) for w in ordered]
    return _Line(
        text=" ".join(w["text"] for w in ordered),
        top=min(float(w["top"]) for w in ordered),
        bottom=max(float(w["bottom"]) for w in ordered),
        x0=min(float(w["x0"]) for w in ordered),
        x1=max(float(w["x1"]) for w in ordered),
        size=max(sizes) if sizes else 0.0,
        page=page_index,
    )


def _containing_table(
    line: _Line, regions: list[tuple[str, tuple[float, float, float, float]]]
) -> str | None:
    centre_y = (line.top + line.bottom) / 2
    centre_x = (line.x0 + line.x1) / 2
    for table_id, (x0, top, x1, bottom) in regions:
        if top <= centre_y <= bottom and x0 <= centre_x <= x1:
            return table_id
    return None


def _table_to_blocks(table, table_id: str, page_index: int, lines: list[_Line]) -> list[_Block]:
    """Preserve table structure rather than flattening it into prose.

    Spec section 31: approval matrices, penalty schedules and expense limits
    live in tables, and a row loses its meaning when its headers are discarded.
    Headers are carried alongside the row rather than woven into a sentence,
    because writing "Tier: 2; Limit: 5000" as prose would be text that does not
    appear anywhere in the document.

    Each row's provenance is the set of real source lines its bounding box
    covers, so the canonical row text is a *recorded transformation* of genuine
    source spans (``table_cell_join``) rather than a free-floating string.
    """

    try:
        rows = table.extract()
        row_objects = list(table.rows)
    except Exception:  # pragma: no cover
        return []
    if not rows:
        return []

    headers = [(cell or "").strip() for cell in rows[0]]
    has_headers = any(headers)

    if not _is_genuine_table(rows):
        return []

    data_rows = list(enumerate(rows))[1:] if has_headers else list(enumerate(rows))

    blocks: list[_Block] = []
    for row_index, row in data_rows:
        cells = [(cell or "").strip() for cell in row]
        if not any(cells):
            continue
        # A pipe-joined cell list keeps every cell value verbatim and adds no
        # words of its own; the separator is structural, not editorial.
        cell_text = " | ".join(cells)

        bbox = None
        if row_index < len(row_objects):
            bbox = getattr(row_objects[row_index], "bbox", None)
        row_lines = _lines_in_bbox(lines, bbox) if bbox else []
        if not row_lines:
            # Coarser but still true: fall back to the whole table's span rather
            # than fabricating a position.
            row_lines = _lines_in_bbox(lines, table.bbox)

        if not row_lines:
            continue

        blocks.append(
            _Block(
                element_type="table_row",
                lines=row_lines,
                table_id=table_id,
                table_headers=headers if has_headers else None,
                cell_text=cell_text,
            )
        )
    return blocks


def _is_genuine_table(rows: list[list[str | None]]) -> bool:
    """Reject bordered text boxes that pdfplumber reports as one-column tables.

    ``find_tables`` works from ruling lines, so any framed callout, shaded
    note, or boxed article body registers as a table. Accepting those is
    actively harmful: each visual line becomes its own "row", which shreds a
    single paragraph into fragments and splits sentences — the exact failure
    spec section 13 prohibits. It is strictly worse than not detecting a table
    at all, because the fragments are then also excluded from paragraph
    reassembly.

    A real policy table (approval matrix, penalty schedule, expense tier) has
    values lined up in at least two columns across at least two rows. A framed
    paragraph has content in exactly one column. That structural difference is
    decidable without heuristics about the text itself.
    """

    if len(rows) < 2:
        return False
    column_fill: Counter[int] = Counter()
    for row in rows:
        for index, cell in enumerate(row):
            if (cell or "").strip():
                column_fill[index] += 1
    populated_columns = [index for index, count in column_fill.items() if count >= 2]
    return len(populated_columns) >= 2


def _lines_in_bbox(
    lines: list[_Line], bbox: tuple[float, float, float, float] | None
) -> list[_Line]:
    if not bbox:
        return []
    x0, top, x1, bottom = bbox
    inside: list[_Line] = []
    for line in lines:
        centre_y = (line.top + line.bottom) / 2
        centre_x = (line.x0 + line.x1) / 2
        if top <= centre_y <= bottom and x0 <= centre_x <= x1:
            inside.append(line)
    return inside


def _modal_body_size(sizes: Counter[float]) -> float:
    """The most common font size is body text; anything larger is structural."""

    if not sizes:
        return 0.0
    return sizes.most_common(1)[0][0]


def _normalize_line(text: str) -> str:
    return _DIGIT_RE.sub("#", " ".join(text.strip().split())).lower()


def _is_edge_line(line: _Line, lines: list[_Line]) -> bool:
    non_empty = [candidate for candidate in lines if candidate.text.strip()]
    if not non_empty:
        return False
    edges = non_empty[:_BOILERPLATE_EDGE_WINDOW] + non_empty[-_BOILERPLATE_EDGE_WINDOW:]
    return any(candidate is line for candidate in edges)


def _detect_boilerplate(page_lines: list[list[_Line]]) -> set[str]:
    """Running headers/footers occupy a fixed page-edge position on most pages.

    Restricting candidates to the exact page edges is what separates a true
    running header from a recurring in-body subheading, which drifts by a line
    or two depending on how long the preceding title is and is real content.
    """

    total_pages = len(page_lines)
    if total_pages < _BOILERPLATE_MIN_PAGES:
        return set()

    counts: Counter[str] = Counter()
    for lines in page_lines:
        non_empty = [line for line in lines if line.text.strip()]
        edges = non_empty[:_BOILERPLATE_EDGE_WINDOW] + non_empty[-_BOILERPLATE_EDGE_WINDOW:]
        seen: set[str] = set()
        for line in edges:
            normalized = _normalize_line(line.text)
            if normalized and normalized not in seen:
                seen.add(normalized)
                counts[normalized] += 1

    threshold = max(_BOILERPLATE_MIN_PAGES, int(total_pages * _BOILERPLATE_MIN_PAGE_FRACTION))
    return {text for text, count in counts.items() if count >= threshold}


_SENTENCE_END = (".", "!", "?", "؟", "。")


def _classify_line(line: _Line, body_size: float) -> ElementType:
    stripped = line.text.strip()
    if not stripped:
        return "other"
    if _LIST_MARKER_RE.match(stripped):
        return "list_item"

    # Font size is strong, direct evidence of a structural heading, so it is
    # trusted on its own.
    larger_font = body_size > 0 and line.size >= body_size + 0.6
    if larger_font and len(stripped) <= 200:
        return "heading"

    # The text-pattern fallback ("ALL CAPS", "Article 12", "7.2.1 ...") is much
    # weaker, so it carries a guard: a line that ends in sentence punctuation is
    # content, not a label. Without this, a short numbered provision such as
    # "2.1 Such jobs shall not be potentially harmful to their health." is
    # classified as a heading, which both loses it as a candidate policy and
    # then mislabels the `section` of every element that follows it. The
    # asymmetry favours guessing "paragraph": a heading demoted to a paragraph
    # is merely a redundant element, whereas a rule promoted to a heading
    # disappears from extraction entirely.
    if len(stripped) <= 90 and not stripped.endswith(_SENTENCE_END):
        if stripped.isupper() or _PROVISION_HEADING_RE.match(stripped):
            return "heading"
    return "paragraph"


def _build_blocks(lines: list[_Line], table_blocks: list[_Block], body_size: float) -> list[_Block]:
    """Group a page's lines into logical blocks, then interleave table rows.

    Blocks break on: a heading, the start of a new list item, entering or
    leaving a table region, and a vertical gap larger than normal leading
    (which is how PDFs express a paragraph break). They deliberately do *not*
    break on line count — that was the old fixed-four-line behaviour the spec
    prohibits (section 12).
    """

    blocks: list[_Block] = []
    current: _Block | None = None
    previous: _Line | None = None
    leading = _typical_leading(lines)

    for line in lines:
        if line.is_boilerplate or not line.text.strip():
            previous = line
            continue
        if line.in_table:
            # Table content is emitted from the structured table extraction
            # instead, so it is not also emitted as loose paragraph text.
            current = None
            previous = line
            continue

        kind = _classify_line(line, body_size)
        gap = (line.top - previous.bottom) if previous is not None else 0.0
        paragraph_break = previous is not None and leading > 0 and gap > leading * 1.6

        starts_new_block = (
            current is None
            or kind == "heading"
            or current.element_type == "heading"
            or kind == "list_item"
            or paragraph_break
        )
        if starts_new_block:
            current = _Block(element_type=kind)
            blocks.append(current)
        current.lines.append(line)
        previous = line

    blocks.extend(table_blocks)
    blocks.sort(key=lambda block: (block.top, block.page))
    return [block for block in blocks if block.lines]


def _typical_leading(lines: list[_Line]) -> float:
    """Median vertical gap between consecutive lines, used to spot paragraph breaks."""

    gaps: list[float] = []
    for previous, current in zip(lines, lines[1:]):
        gap = current.top - previous.bottom
        if 0 <= gap < 60:
            gaps.append(gap)
    if not gaps:
        return 0.0
    gaps.sort()
    return gaps[len(gaps) // 2]


def _join_lines(lines: list[_Line]) -> tuple[str, list[Transformation]]:
    """Concatenate line texts, recording every transformation applied.

    Two joins exist, both deterministic: a single space between lines, and a
    *hyphen-preserving* join when a line ends on a hyphen ("employ-" / "ment"
    becomes "employ-ment", not "employment").

    WHY THE HYPHEN IS KEPT
    ----------------------
    Spec section 30 permits dehyphenation provided it is deterministic, coded
    and recorded. It is not implemented, because it cannot be done correctly:
    a line ending in "non-" followed by "renewal" is indistinguishable from
    "employ-" followed by "ment", yet the first is a real compound hyphen and
    the second is visual line-break hyphenation. Both only ever occur at a
    right margin, so no geometric signal separates them, and most PDFs encode
    both as U+002D rather than using a soft hyphen.

    Removing the hyphen therefore risks producing "nonrenewal" — a token that
    does not appear anywhere in the document — which is precisely what
    INVARIANT 6 exists to prevent. Keeping it produces "employ-ment", which is
    merely ugly: every character still came from the source, a reader and a
    model both parse it correctly, and nothing is fabricated. Given that
    asymmetry, preservation is the only defensible default.

    If dehyphenated text is ever needed for search recall, spec section 30's
    "reversible or source-mapped" clause allows deriving a *separate* field
    from this one. It must not mutate canonical text, which section 28 requires
    to stay stable because every recorded offset indexes into it.
    """

    parts: list[str] = []
    transformations: list[Transformation] = []
    for index, line in enumerate(lines):
        text = line.text
        if index == 0:
            parts.append(text)
            continue
        previous = parts[-1]
        if (
            previous.endswith("-")
            and len(previous) > 1
            and previous[-2].isalpha()
            and text[:1].islower()
        ):
            parts.append(text)
            if "line_break_hyphen_join" not in transformations:
                transformations.append("line_break_hyphen_join")
            continue
        parts.append(" ")
        parts.append(text)
        if "line_join_space" not in transformations:
            transformations.append("line_join_space")
    return "".join(parts), transformations


def _ends_open(text: str) -> bool:
    stripped = text.rstrip()
    if not stripped:
        return False
    return not stripped.endswith(_TERMINAL_PUNCTUATION)


def _continues_previous(previous: _Block, nxt: _Block) -> bool:
    """Decide whether a page-leading block continues the previous page's block.

    Uses several structural signals together rather than one heuristic (spec
    section 9), because any single signal misfires: plenty of legitimate
    paragraphs start lowercase, and plenty of continuations start with a capital
    proper noun.
    """

    if previous.element_type not in ("paragraph", "list_item"):
        return False
    if nxt.element_type not in ("paragraph",):
        return False
    if previous.table_id or nxt.table_id:
        return False

    previous_text = " ".join(line.text for line in previous.lines).strip()
    next_text = " ".join(line.text for line in nxt.lines).strip()
    if not previous_text or not next_text:
        return False

    # Hard boundary signals win outright.
    if _PROVISION_START_RE.match(next_text) or _LIST_MARKER_RE.match(next_text):
        return False

    if not _ends_open(previous_text):
        return False

    first_char = next_text[:1]
    if first_char.islower() or not first_char.isalpha():
        return True
    # An open-ended previous paragraph plus a non-heading, non-provision
    # continuation is still a continuation even when capitalized, since the
    # break mid-sentence is itself strong evidence.
    return True


def _assemble_elements(page_blocks: list[list[_Block]]) -> list[CanonicalElement]:
    """Flatten pages into one ordered element list, merging cross-page blocks."""

    merged: list[_Block] = []
    for blocks in page_blocks:
        for index, block in enumerate(blocks):
            if index == 0 and merged and _continues_previous(merged[-1], block):
                merged[-1].lines.extend(block.lines)
                continue
            merged.append(block)

    elements: list[CanonicalElement] = []
    section: str | None = None
    for order, block in enumerate(merged, start=1):
        if block.cell_text is not None:
            text = block.cell_text
            transformations: list[Transformation] = ["table_cell_join"]
        else:
            text, transformations = _join_lines(block.lines)
        text = text.strip()
        if not text:
            continue
        if len({line.page for line in block.lines}) > 1:
            transformations.append("cross_page_join")

        if block.element_type == "heading":
            section = text

        elements.append(
            CanonicalElement(
                element_id=f"E{len(elements) + 1:06d}",
                element_type=block.element_type,
                logical_order=len(elements),
                text=text,
                section=None if block.element_type == "heading" else section,
                source_fragments=_fragments_for(block.lines),
                transformations=transformations,
                table_id=block.table_id,
                table_headers=block.table_headers,
            )
        )
    return elements


def _fragments_for(lines: list[_Line]) -> list[SourceFragment]:
    """One fragment per contiguous run of lines on the same page.

    Adjacent lines on one page collapse into a single span so a cross-page
    element reports exactly one fragment per page it touches, which is the
    shape spec section 25 specifies and section 27 validates against.
    """

    fragments: list[SourceFragment] = []
    for line in lines:
        if fragments and fragments[-1].page == line.page and fragments[-1].end_offset <= line.start_offset:
            previous = fragments[-1]
            fragments[-1] = SourceFragment(
                page=previous.page,
                start_offset=previous.start_offset,
                end_offset=line.end_offset,
                text=previous.text + "\n" + line.text
                if line.start_offset == previous.end_offset + 1
                else previous.text + " " + line.text,
            )
        else:
            fragments.append(line.fragment)
    return fragments

# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def ingest_docx(storage_path: str | Path, document_id: str = "") -> CanonicalDocument:
    """Parse a DOCX into the same canonical representation as a PDF.

    DOCX has no page structure — pagination is decided by the renderer, not
    stored in the file — so the whole document is modelled as one logical page
    whose ``raw_text`` is every block joined in document order. Offsets index
    into that string, so provenance stays exactly as verifiable as it is for a
    PDF, and downstream consumers need no special case.

    The previous DOCX path turned each table row into prose of the form
    ``"Tier: 2; Limit: 5000"``. That text appears nowhere in the document, so
    any passage "quoted" from it was fabricated by definition — the very thing
    INVARIANT 6 exists to prevent. Rows now keep their cell values verbatim and
    carry their headers as structured data instead.
    """

    path = Path(storage_path)
    try:
        document = DocxDocument(str(path))
    except Exception as exc:
        raise IngestionError(f"cannot open DOCX {path.name}: {exc}") from exc

    diagnostics: list[IngestionDiagnostic] = []
    raw_parts: list[str] = []
    offset = 0
    elements: list[CanonicalElement] = []
    section: str | None = None
    table_index = 0

    def _emit(
        text: str,
        element_type: ElementType,
        table_id: str | None = None,
        headers: list[str] | None = None,
        transformations: list[Transformation] | None = None,
        raw_line: str | None = None,
    ) -> None:
        nonlocal offset, section
        line = raw_line if raw_line is not None else text
        if not line.strip():
            return
        start = offset
        end = start + len(line)
        raw_parts.append(line)
        offset = end + 1  # the newline joining blocks

        if element_type == "heading":
            section = text
        elements.append(
            CanonicalElement(
                element_id=f"E{len(elements) + 1:06d}",
                element_type=element_type,
                logical_order=len(elements),
                text=text,
                section=None if element_type == "heading" else section,
                source_fragments=[
                    SourceFragment(page=1, start_offset=start, end_offset=end, text=line)
                ],
                transformations=transformations or [],
                table_id=table_id,
                table_headers=headers,
            )
        )

    for block in _iter_docx_blocks(document):
        if isinstance(block, DocxParagraph):
            text = block.text.strip()
            if not text:
                continue
            _emit(text, _classify_docx_paragraph(block, text))
        elif isinstance(block, DocxTable):
            table_index += 1
            table_id = f"t{table_index}"
            try:
                rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
            except Exception as exc:
                diagnostics.append(
                    IngestionDiagnostic(
                        code="table_parse_failed",
                        severity="warning",
                        detail=f"table {table_id}: {exc}",
                    )
                )
                continue
            if not rows:
                continue
            headers = rows[0]
            has_headers = any(headers)
            genuine = _is_genuine_table(rows)
            data_rows = rows[1:] if (has_headers and genuine) else rows
            for row in data_rows:
                if not any(row):
                    continue
                if genuine:
                    _emit(
                        " | ".join(row),
                        "table_row",
                        table_id=table_id,
                        headers=headers if has_headers else None,
                        transformations=["table_cell_join"],
                    )
                else:
                    # A single-column "table" is a layout device, so its cells
                    # are ordinary paragraphs.
                    for cell in row:
                        if cell.strip():
                            _emit(cell, "paragraph")

    canonical = CanonicalDocument(
        document_id=document_id or path.stem,
        page_count=1,
        pages=[CanonicalPage(page=1, raw_text="\n".join(raw_parts))],
        elements=elements,
        parser=DOCX_PARSER_NAME,
        diagnostics=diagnostics,
    )
    _append_document_diagnostics(canonical)
    return canonical


def _classify_docx_paragraph(block: DocxParagraph, text: str) -> ElementType:
    """Prefer the document's own declared style over guessing from the text."""

    try:
        style_name = (block.style.name if block.style else "") or ""
    except Exception:
        style_name = ""
    lowered = style_name.lower()
    if lowered.startswith("heading") or lowered in ("title", "subtitle"):
        return "heading"
    if "list" in lowered:
        return "list_item"
    if _LIST_MARKER_RE.match(text):
        return "list_item"
    if len(text) <= 90 and not text.endswith(_SENTENCE_END):
        if text.isupper() or _PROVISION_HEADING_RE.match(text):
            return "heading"
    return "paragraph"


def _iter_docx_blocks(document):
    """Yield paragraphs and tables in true document order.

    python-docx exposes ``paragraphs`` and ``tables`` as separate collections
    with no interleaving, so walking the XML body is the only way to preserve
    the order a reader sees — and order is what makes a table row attributable
    to the section heading above it.
    """

    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield DocxParagraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield DocxTable(child, document)
