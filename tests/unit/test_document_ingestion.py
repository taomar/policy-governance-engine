"""Boundary tests for canonical document ingestion (spec sections 47-53).

These target the failures that are *invisible downstream*: text that reads
plausibly but was assembled wrongly. A cut paragraph, a welded two-column line,
or a shredded table all produce output that looks like ordinary prose, so
nothing later in the pipeline can flag them — they have to be caught here.
"""
from __future__ import annotations

from pathlib import Path

import pytest

from policy_platform.contracts.canonical_document import (
    CanonicalDocument,
    CanonicalElement,
    CanonicalPage,
    SourceFragment,
)
from policy_platform.infrastructure.ingestion import document_ingestion as ingestion
from tests.corpus import tracked_document, uploaded_document
from policy_platform.infrastructure.ingestion.document_ingestion import (
    IngestionError,
    _Block,
    _classify_line,
    _continues_previous,
    _detect_columns,
    _is_genuine_table,
    _join_lines,
    _Line,
)

DOCUMENTS = Path(__file__).resolve().parents[2] / "data" / "documents"
#: Pinned by their full upload name, these two resolved nowhere: `data/documents/`
#: is gitignored and its filenames carry the upload UUID, so both classes below
#: skipped in every checkout and had never run. `test_table_headers_are_preserved`
#: guards the header-row defect fixed in this same file.
#: `TestRealPdf` asserts invariants that hold of any real document, so it is
#: pointed at a committed one. It was previously pinned to an upload named
#: `...saudilaborlaw.pdf` that is present in no checkout, and skipped silently
#: in all of them.
REAL_PDF = "HR-Guide-Policy-and-Procedure-Template.pdf"
HARDWARE_DOCX = "Workplace-Hardware-Provisioning-Policy-v3.3.docx"


def _line(text: str, *, top=0.0, x0=0.0, x1=100.0, size=10.0, page=1) -> _Line:
    return _Line(
        text=text,
        top=top,
        bottom=top + 10,
        x0=x0,
        x1=x1,
        size=size,
        page=page,
        start_offset=0,
        end_offset=len(text),
    )


class TestLineJoining:
    """Spec section 10: only characters from source fragments may participate."""

    def test_joins_lines_with_a_single_space(self):
        text, transformations = _join_lines([_line("If an employee is absent"), _line("for thirty days")])
        assert text == "If an employee is absent for thirty days"
        assert "line_join_space" in transformations

    def test_preserves_hyphen_rather_than_guessing_a_word(self):
        # "non-" + "renewal" must not become "nonrenewal", a token that appears
        # nowhere in the source.
        text, transformations = _join_lines([_line("notice of non-"), _line("renewal shall be given")])
        assert text == "notice of non-renewal shall be given"
        assert "line_break_hyphen_join" in transformations

    def test_never_inserts_a_word(self):
        text, _ = _join_lines([_line("The employee must obtain"), _line("approval prior to travel.")])
        assert text == "The employee must obtain approval prior to travel."
        assert "manager" not in text

    def test_single_line_needs_no_transformation(self):
        text, transformations = _join_lines([_line("A standalone sentence.")])
        assert text == "A standalone sentence."
        assert transformations == []


class TestClassification:
    def test_provision_label_is_a_heading(self):
        assert _classify_line(_line("Article 12"), 10.0) == "heading"
        assert _classify_line(_line("Article (81)"), 10.0) == "heading"
        assert _classify_line(_line("Chapter 1: Definitions"), 10.0) == "heading"

    def test_provision_citation_inside_a_sentence_is_not_a_heading(self):
        # This is a real policy clause that merely begins with a citation.
        # Promoting it to a heading loses it from extraction entirely and
        # mislabels the section of everything that follows.
        line = _line("Article (81) of the Labor Law during the training term or within")
        assert _classify_line(line, 10.0) == "paragraph"

    def test_numbered_sentence_is_not_a_heading(self):
        line = _line("2.1 Such jobs shall not be potentially harmful to their health.")
        assert _classify_line(line, 10.0) == "paragraph"

    def test_larger_font_is_a_heading(self):
        assert _classify_line(_line("Scope And Purpose", size=14.0), 10.0) == "heading"

    def test_bullet_is_a_list_item(self):
        assert _classify_line(_line("1. Employees shall submit a request."), 10.0) == "list_item"
        assert _classify_line(_line("(a) written notice"), 10.0) == "list_item"


class TestGenuineTable:
    """Spec section 13: a chunk boundary must never split a sentence."""

    def test_multi_column_table_is_genuine(self):
        rows = [["Tier", "Limit"], ["1", "5000"], ["2", "10000"]]
        assert _is_genuine_table(rows) is True

    def test_bordered_paragraph_is_not_a_table(self):
        # A framed callout registers as a one-column table; accepting it splits
        # one paragraph into four "rows".
        rows = [
            ["", "Article (6)", ""],
            ["", "Incidental workers shall be subject to the provisions of duties,", ""],
            ["", "disciplinary rules, the maximum working hours, and daily rest.", ""],
        ]
        assert _is_genuine_table(rows) is False

    def test_single_row_is_not_a_table(self):
        assert _is_genuine_table([["Tier", "Limit"]]) is False


class TestColumnDetection:
    def test_single_column_prose_has_no_gutter(self):
        lines = [_line(f"line {i}", top=i * 12, x0=50, x1=550) for i in range(20)]
        assert _detect_columns(lines, 600.0) == []

    def test_two_column_layout_is_detected(self):
        left = [_line(f"left {i}", top=i * 12, x0=50, x1=270) for i in range(12)]
        right = [_line(f"right {i}", top=i * 12, x0=330, x1=550) for i in range(12)]
        columns = _detect_columns(left + right, 600.0)
        assert len(columns) == 2

    def test_ragged_right_margin_is_not_a_gutter(self):
        lines = [_line(f"line {i}", top=i * 12, x0=50, x1=400 + (i % 5) * 20) for i in range(20)]
        assert _detect_columns(lines, 600.0) == []

    def test_too_few_lines_yields_no_columns(self):
        lines = [_line("a", top=0, x0=50, x1=100), _line("b", top=12, x0=400, x1=450)]
        assert _detect_columns(lines, 600.0) == []


class TestCrossPageContinuation:
    def _block(self, kind, text, page=1):
        return _Block(element_type=kind, lines=[_line(text, page=page)])

    def test_open_sentence_continues_onto_next_page(self):
        previous = self._block("paragraph", "If an employee is absent from work for more than", page=1)
        nxt = self._block("paragraph", "thirty consecutive days, the employer may terminate.", page=2)
        assert _continues_previous(previous, nxt) is True

    def test_completed_sentence_does_not_continue(self):
        previous = self._block("paragraph", "The employer shall keep records.", page=1)
        nxt = self._block("paragraph", "Wages shall be paid in the national currency.", page=2)
        assert _continues_previous(previous, nxt) is False

    def test_new_article_is_a_hard_boundary(self):
        previous = self._block("paragraph", "the employer may terminate the contract without", page=1)
        nxt = self._block("paragraph", "Article 80 The employer may terminate without award", page=2)
        assert _continues_previous(previous, nxt) is False

    def test_new_list_item_is_a_hard_boundary(self):
        previous = self._block("paragraph", "the following shall apply to any worker who", page=1)
        nxt = self._block("paragraph", "1. fails to report for duty", page=2)
        assert _continues_previous(previous, nxt) is False

    def test_heading_never_continues_a_paragraph(self):
        previous = self._block("paragraph", "the employer may terminate the contract without", page=1)
        nxt = self._block("heading", "Article 81", page=2)
        assert _continues_previous(previous, nxt) is False


class TestFragmentVerification:
    def test_resolvable_offsets_pass(self):
        page = CanonicalPage(page=1, raw_text="The employer shall keep records.")
        element = CanonicalElement(
            element_id="E000001",
            element_type="paragraph",
            logical_order=0,
            text="The employer",
            source_fragments=[SourceFragment(page=1, start_offset=0, end_offset=12, text="The employer")],
        )
        document = CanonicalDocument(
            document_id="d", page_count=1, pages=[page], elements=[element], parser="test"
        )
        assert document.verify_fragments() == []

    def test_unresolvable_offsets_are_reported(self):
        page = CanonicalPage(page=1, raw_text="The employer shall keep records.")
        element = CanonicalElement(
            element_id="E000001",
            element_type="paragraph",
            logical_order=0,
            text="something else",
            source_fragments=[SourceFragment(page=1, start_offset=0, end_offset=12, text="Wrong text!!")],
        )
        document = CanonicalDocument(
            document_id="d", page_count=1, pages=[page], elements=[element], parser="test"
        )
        assert len(document.verify_fragments()) == 1


class TestUnsupportedInput:
    def test_unknown_type_is_rejected_clearly(self, tmp_path):
        target = tmp_path / "notes.txt"
        target.write_text("hello")
        with pytest.raises(IngestionError):
            ingestion.ingest_document(target)

    def test_corrupt_pdf_raises_rather_than_returning_empty(self, tmp_path):
        target = tmp_path / "broken.pdf"
        target.write_bytes(b"%PDF-1.4 this is not a real pdf")
        with pytest.raises(IngestionError):
            ingestion.ingest_document(target)


class TestRealPdf:
    """Spec sections 51-53 against a real document.

    Resolved by the stable tail of the upload name rather than the whole of it,
    and absent means failed rather than skipped -- see `tests/corpus.py`.
    """

    @pytest.fixture(scope="class")
    def source(self):
        return tracked_document(REAL_PDF)

    @pytest.fixture(scope="class")
    def document(self, source):
        return ingestion.ingest_pdf(source, "real-pdf")

    def test_every_page_is_ingested(self, document, source):
        # INVARIANT 1. The count is read from the file rather than written here,
        # so re-issuing the document does not turn this into a false alarm.
        import pdfplumber

        with pdfplumber.open(source) as pdf:
            expected = len(pdf.pages)
        assert expected > 1, "a single-page document cannot exercise this invariant"
        assert document.page_count == expected
        assert len(document.pages) == expected

    def test_every_fragment_resolves(self, document):
        # INVARIANT 4 and 5.
        assert document.verify_fragments() == []

    def test_elements_form_a_total_order(self, document):
        # INVARIANT 2.
        orders = [element.logical_order for element in document.elements]
        assert orders == sorted(orders)
        assert len(set(orders)) == len(orders)

    def test_paragraphs_are_reconnected_across_pages(self, document):
        # INVARIANT 3: a page boundary is not a semantic boundary.
        assert any(element.spans_pages for element in document.elements)

    def test_no_word_is_invented(self, document):
        # INVARIANT 6, checked over the whole document rather than a sample.
        import re
        from collections import Counter

        def words(value: str) -> list[str]:
            return re.findall(r"\w+", value.lower())

        for element in document.elements:
            if element.table_id:
                continue  # table text is a recorded cell join, not a line join
            available: Counter[str] = Counter()
            for fragment in element.source_fragments:
                available.update(words(fragment.text))
            for word in words(element.text):
                assert available[word] > 0, f"{element.element_id} invented {word!r}"
                available[word] -= 1

    def test_coverage_is_near_total(self, document):
        # INVARIANT 8: the only characters that may go missing are the running
        # headers deliberately removed from the logical flow.
        raw = sum(len("".join(page.raw_text.split())) for page in document.pages)
        kept = sum(len("".join(element.text.split())) for element in document.elements)
        assert kept / raw > 0.95

    def test_ingestion_is_deterministic(self, source):
        # Spec section 45: identical bytes produce identical ids and offsets.
        first = ingestion.ingest_pdf(source, "real-pdf")
        second = ingestion.ingest_pdf(source, "real-pdf")
        assert [e.element_id for e in first.elements] == [e.element_id for e in second.elements]
        assert [e.text for e in first.elements] == [e.text for e in second.elements]
        assert [
            (f.page, f.start_offset, f.end_offset)
            for e in first.elements
            for f in e.source_fragments
        ] == [
            (f.page, f.start_offset, f.end_offset)
            for e in second.elements
            for f in e.source_fragments
        ]

    def test_sections_are_attributed(self, document):
        paragraphs = [e for e in document.elements if e.element_type == "paragraph"]
        attributed = [e for e in paragraphs if e.section]
        assert len(attributed) / len(paragraphs) > 0.5


class TestDocxHeaderRowIsEvidencedToo:
    """The standard `fa27428` established held on one parser only.

    A DOCX table's row 0 was read as the header whenever any of its cells held
    text, then dropped with `rows[1:]` -- the same content-loss defect, on the
    other path, with no diagnostic anywhere in it, so a reviewer had no way to
    discover that a row had been consumed or why.

    These grids are synthetic and carry no words from any document. They vary
    only in form, which is the only thing the decision is allowed to read.
    """

    @staticmethod
    def _docx(tmp_path, grid: list[list[str]], name: str = "g.docx"):
        from docx import Document as DocxDocument

        document = DocxDocument()
        table = document.add_table(rows=len(grid), cols=len(grid[0]))
        for row_index, row in enumerate(grid):
            for column_index, cell in enumerate(row):
                table.cell(row_index, column_index).text = cell
        path = tmp_path / name
        document.save(str(path))
        return ingestion.ingest_docx(str(path), "g")

    @staticmethod
    def _rows(document):
        return [e for e in document.elements if e.element_type == "table_row"]

    def test_a_row_whose_cells_recur_below_is_kept_as_content(self, tmp_path):
        """Absence before, presence after: row 0 survives instead of vanishing."""

        grid = [["repeated", "shared"], ["repeated", "shared"], ["other", "shared"]]
        document = self._docx(tmp_path, grid)

        assert " | ".join(grid[0]) in {row.text for row in self._rows(document)}

    def test_an_unevidenced_grid_carries_no_labels_at_all(self, tmp_path):
        """Absent, not empty. `None` says no row stated labels; `[]` would not."""

        grid = [["repeated", "shared"], ["repeated", "shared"], ["other", "shared"]]
        document = self._docx(tmp_path, grid)

        for row in self._rows(document):
            assert row.table_headers is None

    def test_a_docx_reviewer_can_discover_that_no_row_stated_labels(self, tmp_path):
        """The diagnostic had exactly one call site, inside the PDF function."""

        grid = [["repeated", "shared"], ["repeated", "shared"], ["other", "shared"]]
        document = self._docx(tmp_path, grid)

        reported = [
            d
            for d in document.diagnostics
            if d.code == "table_header_row_not_identified"
        ]
        assert reported, "a DOCX table's headerless verdict reached no reviewer"
        assert "recur further down" in reported[0].detail
        assert reported[0].severity == "warning"

    def test_a_banner_row_is_not_read_as_labels(self, tmp_path):
        grid = [["banner", "", ""], ["a", "b", "c"], ["d", "e", "f"]]
        document = self._docx(tmp_path, grid)

        assert " | ".join(grid[0]) in {row.text for row in self._rows(document)}

    def test_an_evidenced_header_is_still_consumed_and_still_carried(self, tmp_path):
        """The reluctance must not become refusal: real headers still work."""

        grid = [["alpha", "beta"], ["a1", "b1"], ["a2", "b2"]]
        document = self._docx(tmp_path, grid)

        rows = self._rows(document)
        assert " | ".join(grid[0]) not in {row.text for row in rows}
        for row in rows:
            assert row.table_headers == grid[0]
        assert not [
            d
            for d in document.diagnostics
            if d.code == "table_header_row_not_identified"
        ]


def test_both_parsers_ask_the_same_question_in_one_place():
    """A second copy of this decision is how it came to hold on one path only.

    The verdict is reached once, in `_column_labels_for`, so the two parsers
    cannot drift again: a new call to the underlying test elsewhere in the
    module means someone has begun forking it.
    """

    import ast

    source = Path(ingestion.__file__).read_text(encoding="utf-8")
    tree = ast.parse(source)

    holders = {
        node.name
        for node in ast.walk(tree)
        if isinstance(node, ast.FunctionDef)
        and any(
            isinstance(call.func, ast.Name)
            and call.func.id == "_row_states_column_labels"
            for call in ast.walk(node)
            if isinstance(call, ast.Call)
        )
    }
    assert holders == {"_column_labels_for"}


class TestRealDocx:
    @pytest.fixture(scope="class")
    def document(self):
        return ingestion.ingest_docx(tracked_document(HARDWARE_DOCX), "hardware")

    def test_fragments_resolve(self, document):
        assert document.verify_fragments() == []

    def test_table_rows_keep_cell_values_verbatim(self, document):
        rows = [e for e in document.elements if e.element_type == "table_row"]
        assert rows, "expected the approval-tier tables to survive ingestion"
        raw = document.page_text(1)
        for row in rows:
            # Each cell value must appear in the source; the pipe separator is
            # structural, never editorial.
            for cell in row.text.split(" | "):
                if cell.strip():
                    assert cell.strip() in raw

    def test_table_headers_are_preserved_not_flattened(self, document):
        rows = [e for e in document.elements if e.element_type == "table_row" and e.table_headers]
        assert rows, "expected at least one table with headers"
        for row in rows:
            assert "; " not in row.text or " | " in row.text


class TestHeaderRowIsEvidencedNotAssumed:
    """`page.find_tables()` is a per-page API: it has no cross-page concept, so a
    table running onto the next page is re-discovered there as a fresh grid whose
    row 0 is *content*, not labels.

    Ingestion used to treat row 0 as the header row whenever it held any non-empty
    cell, then drop it with `rows[1:]`. On a continuation page that deleted a row
    of the schedule outright and stamped its text onto the surviving rows as their
    column headers -- so the field was populated on every row and correct on none.

    These cases are written on cell *form* only. Nothing keys on a heading, a
    numbering scheme or a layout, so they hold for any document.
    """

    def test_short_distinct_labels_are_accepted_as_a_header_row(self):
        rows = [
            ["Grade", "Basic Salary", "Housing"],
            ["A", "10,000", "2,000"],
            ["B", "8,000", "1,600"],
        ]
        accepted, reason = ingestion._row_states_column_labels(rows)
        assert accepted, reason
        assert reason == ""

    def test_a_continuation_pages_first_row_is_content_not_a_header(self):
        # The distinguishing mark is length: a label names a column, a cell of a
        # schedule states a provision. Nothing here depends on what it says.
        rows = [
            ["7", "Absence without written permission for seven to ten days "
                  "within a single contract year.", "Final written warning"],
            ["8", "Departing early without permission.", "Verbal warning"],
        ]
        accepted, reason = ingestion._row_states_column_labels(rows)
        assert not accepted
        assert "longer than" in reason

    def test_a_row_spanned_by_one_banner_cell_is_not_a_header(self):
        # A merged family banner paints into one cell of a wide grid. It labels
        # nothing; it announces the block beneath it.
        rows = [
            ["Working hours", "", "", "", ""],
            ["1", "Late arrival", "Warning", "5%", "10%"],
        ]
        accepted, reason = ingestion._row_states_column_labels(rows)
        assert not accepted
        assert "1 of 5" in reason

    def test_a_row_whose_values_recur_below_is_not_a_header(self):
        # Labels are distinct from the values under them; values repeat. A grid
        # that opens mid-block can look label-shaped by length alone, so this
        # catches the short-celled continuation the length rule cannot.
        rows = [
            ["Warning", "5%", "10%"],
            ["Warning", "5%", "10%"],
            ["Dismissal", "5%", "10%"],
        ]
        accepted, reason = ingestion._row_states_column_labels(rows)
        assert not accepted
        assert "recur further down" in reason

    def test_a_single_celled_row_is_not_a_header(self):
        rows = [["Section 4"], ["Some provision"]]
        accepted, reason = ingestion._row_states_column_labels(rows)
        assert not accepted
        assert "1 of 1" in reason

    def test_an_empty_grid_states_no_labels(self):
        assert ingestion._row_states_column_labels([]) == (False, "the grid has no rows")

    def test_a_two_column_grid_cannot_evidence_which_way_its_labels_run(self):
        """Form alone does not separate a header row from a key/value first row.

        A grid two columns wide is symmetric under transposition: "the labels
        run across the top" and "the keys run down the left side" draw the same
        picture. Both grids below have a first row of two filled cells, each
        short enough to be a label, neither recurring beneath -- every property
        this test is allowed to look at. One opens with a header; the other
        opens with an ordinary pair. Nothing in their shape says which.

        Separating them needs a judgement about what the words mean, and the
        contract of `_row_states_column_labels` is that it keys on cell form and
        never on cell content. So the two verdicts must agree. If this test ever
        fails, a content judgement has been introduced -- which is a decision to
        take deliberately, in one language at a time, not a tidy-up.
        """

        labels_across_the_top = [
            ["Abbreviation", "Meaning"],
            ["ABC", "a phrase the abbreviation stands for"],
            ["DEF", "another phrase it stands for"],
        ]
        keys_down_the_side = [
            ["Reference", "a value recorded against the reference"],
            ["Owner", "a value recorded against the owner"],
            ["Status", "a value recorded against the status"],
        ]

        assert ingestion._row_states_column_labels(labels_across_the_top) == (
            ingestion._row_states_column_labels(keys_down_the_side)
        )


class _StubRow:
    """One row's geometry, with a distinct x-slice per column.

    The cells used to share a single bbox spanning the whole row. That was
    enough while the only thing read from a row was `bbox`, for provenance —
    but a real pdfplumber row gives each column its own x-range, and a cell
    that spans several columns is how a merged banner is expressed. A stub
    where every cell covers every column says "every header cell is a banner
    over all the others", which is not a table any parser would produce.
    """

    def __init__(self, top: float, cell_count: int):
        self.bbox = (0.0, top, 100.0, top + 10.0)
        width = 100.0 / cell_count if cell_count else 100.0
        self.cells = [
            (i * width, top, (i + 1) * width, top + 10.0) for i in range(cell_count)
        ]


class _StubPage:
    width = 200.0
    height = 800.0


class _StubTable:
    """The narrowest surface `_table_to_blocks` touches, so the case can be run
    without a PDF. `data/documents/` is gitignored, so a corpus-backed guard
    would skip silently in a clean checkout -- which is the failure this
    repository has already absorbed five times."""

    page = _StubPage()

    def __init__(self, grid):
        self._grid = grid
        self.rows = [_StubRow(i * 10.0, len(r)) for i, r in enumerate(grid)]
        self.bbox = (0.0, 0.0, 100.0, len(grid) * 10.0)

    def extract(self):
        return self._grid


def _blocks_for(grid):
    table = _StubTable(grid)
    lines = [
        _line(" ".join(c for c in row if c), top=i * 10.0)
        for i, row in enumerate(grid)
    ]
    return ingestion._table_to_blocks(table, "p22-t1", 22, lines)


class TestContinuationRowSurvivesAndSaysSo:
    """Two obligations. The row must not be lost -- content loss is the defect.
    And a table whose header row could not be identified must be *discoverable*,
    because a reviewer judging coverage cannot see an absence that nothing reports.
    """

    CONTINUATION = [
        ["7", "Absence without written permission for seven to ten days within a contract year.", "Final warning"],
        ["8", "Departing from work fifteen minutes early without a valid reason.", "Verbal warning"],
        ["9", "Sleeping while on duty in a manner that endangers others.", "Suspension"],
    ]

    HEADED = [
        ["No.", "Violation", "Penalty"],
        ["1", "Late arrival without permission.", "Warning"],
        ["2", "Leaving the site without notice.", "Warning"],
    ]

    def test_the_first_row_of_a_continuation_page_is_emitted_as_content(self):
        blocks, _ = _blocks_for(self.CONTINUATION)
        emitted = "\n".join(b.cell_text or "" for b in blocks)
        assert "Absence without written permission" in emitted, (
            "the continuation page's leading row was consumed as a header and lost"
        )
        assert len(blocks) == 3

    def test_no_row_is_labelled_with_another_rows_text(self):
        blocks, _ = _blocks_for(self.CONTINUATION)
        assert [b.table_headers for b in blocks] == [None, None, None], (
            "an unidentified header row must leave the field empty; a wrong value "
            "invites trust where an empty one invites a question"
        )

    def test_the_failure_is_reported_rather_than_passed_over(self):
        _, diagnostics = _blocks_for(self.CONTINUATION)
        assert len(diagnostics) == 1
        diagnostic = diagnostics[0]
        assert diagnostic.code == "table_header_row_not_identified"
        assert diagnostic.severity == "warning"
        assert diagnostic.page == 22
        assert "p22-t1" in diagnostic.detail
        assert "longer than" in diagnostic.detail, (
            "the diagnostic must name what disqualified the row, not merely that "
            "something did"
        )

    def test_a_real_header_row_is_still_lifted_and_reports_nothing(self):
        blocks, diagnostics = _blocks_for(self.HEADED)
        assert diagnostics == []
        assert len(blocks) == 2, "the label row is carried structurally, not emitted"
        assert all(b.table_headers == ["No.", "Violation", "Penalty"] for b in blocks)
        assert "Late arrival" in (blocks[0].cell_text or "")

    def test_cells_are_still_joined_by_the_separator_the_safety_gate_reads(self):
        # `formulation_mapping.states_a_flattened_row` keys on " | " to recognise a
        # flattened table row. Recovering rows makes that marker appear on *more*
        # records, never fewer, so the gate holds wider rather than narrower.
        blocks, _ = _blocks_for(self.CONTINUATION)
        assert all(" | " in (b.cell_text or "") for b in blocks)


class TestWithinPageContinuation:
    """Merging used to be gated on `index == 0`, so only a page-leading block
    could ever be joined and a sentence cut mid-page stayed cut.

    Two clauses in AD-103 begin mid-sentence for that reason, and each produced
    a duplicate rule: the formulator reconstructs the governing sentence from
    inherited context for the orphaned half, remaking a rule the preceding
    clause already produced. Two cuts, two duplicate pairs, exact
    correspondence.
    """

    def _block(self, kind, text, page=1):
        return _Block(element_type=kind, lines=[_line(text, page=page)])

    def test_a_bracket_continues_within_a_page(self):
        """The real cut: "...one employee of the married couple" | "(husband
        and wife). In the case of...". Both on page 2."""

        previous = self._block(
            "paragraph", "The housing allowance is limited to one employee of the married couple", page=2
        )
        nxt = self._block(
            "paragraph", "(husband and wife). In the case of a married couple are employed by FBSU", page=2
        )
        assert _continues_previous(previous, nxt, same_page=True) is True

    def test_a_lowercase_start_continues_within_a_page(self):
        """The other real cut: "...is calculated as twice" | "the monthly basic
        salary up to a maximum of:"."""

        previous = self._block(
            "paragraph", "The housing allowance per calendar year (12 months) is calculated as twice", page=2
        )
        nxt = self._block("paragraph", "the monthly basic salary up to a maximum of:", page=2)
        assert _continues_previous(previous, nxt, same_page=True) is True

    def test_a_capitalised_block_does_not_continue_within_a_page(self):
        """Within a page a block break carries layout meaning — spacing, a font
        change, a new column — so a capitalised block after an unpunctuated one
        is commonly a genuine new paragraph. Across a page break the same pair
        does continue, because a sentence does not normally end at a page
        boundary without punctuation."""

        previous = self._block("paragraph", "The employer shall keep records of", page=2)
        nxt = self._block("paragraph", "Employees are entitled to annual leave.", page=2)
        assert _continues_previous(previous, nxt, same_page=True) is False
        assert _continues_previous(previous, nxt, same_page=False) is True

    def test_a_completed_sentence_never_continues(self):
        previous = self._block("paragraph", "The housing allowance is paid monthly.", page=2)
        nxt = self._block("paragraph", "(husband and wife) are both eligible.", page=2)
        assert _continues_previous(previous, nxt, same_page=True) is False

    def test_a_provision_number_is_still_a_hard_boundary(self):
        previous = self._block("paragraph", "The housing allowance is limited to one employee of", page=2)
        nxt = self._block("paragraph", "3.4.2. Transportation allowance is paid separately.", page=2)
        assert _continues_previous(previous, nxt, same_page=True) is False

    def test_a_list_marker_is_still_a_hard_boundary(self):
        previous = self._block("paragraph", "The allowance is limited to one employee of", page=2)
        nxt = self._block("paragraph", "1. the married couple", page=2)
        assert _continues_previous(previous, nxt, same_page=True) is False
